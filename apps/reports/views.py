import matplotlib
matplotlib.use('Agg')  # Set non-interactive backend to prevent GUI errors

from django.shortcuts import render
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML, CSS
from risk.models import Risk, RiskRegister, Control, KRI, RiskAssessment
from django.db.models import Count, Q, Avg, Sum
from django.db import models
import tempfile
import matplotlib.pyplot as plt
import io
import base64
from datetime import datetime
from django.db.models.functions import TruncMonth
from audit.models import AuditWorkplan, Engagement, Issue, Approval
from audit.models.recommendation import Recommendation
from django.db.models.functions import TruncYear
from legal.models import LegalCase, LegalTask, LegalDocument, LegalParty, LegalArchive
from compliance.models import ComplianceRequirement, ComplianceFramework, PolicyDocument, ComplianceObligation, ComplianceEvidence
from datetime import timedelta
from contracts.models import Contract, Party, ContractMilestone, ContractType
from django.utils import timezone
from risk.models import Objective

def _docx_start_document(org, title, generation_timestamp):
    """Create a python-docx Document with standard header and metadata."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Organization logo (best-effort)
    if getattr(org, 'logo', None) and getattr(org.logo, 'path', None):
        try:
            doc.add_picture(org.logo.path, width=Inches(1.8))
        except Exception:
            pass

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata table
    meta = doc.add_table(rows=2, cols=2)
    meta.style = 'Table Grid'
    meta.cell(0, 0).text = 'Organization'
    meta.cell(0, 1).text = getattr(org, 'name', 'Organization')
    meta.cell(1, 0).text = 'Generated'
    meta.cell(1, 1).text = generation_timestamp
    return doc

def _docx_http_response(doc, filename_prefix, org):
    """Return the given Document as an HTTP response download."""
    from django.http import HttpResponse
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{org.code}_{filename_prefix}.docx"'
    doc.save(response)
    return response

def _html_to_text(html_string: str) -> str:
    """Convert rich HTML from editors into readable plain text for DOCX."""
    if not html_string:
        return ''
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(str(html_string), 'html.parser')
        # Ensure line breaks for <br> and paragraphs
        for br in soup.find_all('br'):
            br.replace_with('\n')
        # Separate block elements with newlines
        for block in soup.find_all(['p', 'div', 'li']):
            if block.string is None:
                # Add newline after block content if not already
                block.append('\n')
        text = soup.get_text()
        # Normalize excessive blank lines
        lines = [l.rstrip() for l in text.splitlines()]
        # Remove trailing empty lines
        while lines and not lines[-1]:
            lines.pop()
        return '\n'.join(lines)
    except Exception:
        # Fallback: plain string coercion
        return str(html_string)

def _docx_add_heading(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def _docx_add_html_block(doc, html_string: str):
    text = _html_to_text(html_string)
    if not text:
        return
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)

def risk_report_pdf(request):
    org = request.tenant
    risks = Risk.objects.filter(organization=org)
    # Apply filters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    status = request.GET.get('status')
    if start_date:
        risks = risks.filter(date_identified__gte=start_date)
    if end_date:
        risks = risks.filter(date_identified__lte=end_date)
    if status:
        risks = risks.filter(status=status)
    html_string = render_to_string('reports/risk_report.html', {
        'organization': org,
        'risks': risks,
        'title': 'Risk Register Report',
        'description': f'This report summarizes all risks for your organization. Filters: {start_date or "Any"} to {end_date or "Any"}, Status: {status or "All"}.',
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_risk_report.pdf"'
    return response

def risk_register_summary_pdf(request):
    org = request.tenant
    risks = Risk.objects.filter(organization=org)
    category = request.GET.get('category')
    owner = request.GET.get('owner')
    status = request.GET.get('status')
    register = request.GET.get('register')
    if category:
        risks = risks.filter(category=category)
    if owner:
        risks = risks.filter(risk_owner__icontains=owner)
    if status:
        risks = risks.filter(status=status)
    if register:
        risks = risks.filter(risk_register__register_name__icontains=register)
    summary = {
        'by_status': risks.values('status').annotate(count=Count('id')),
        'by_category': risks.values('category').annotate(count=Count('id')),
        'by_owner': risks.values('risk_owner').annotate(count=Count('id')),
        'by_register': risks.values('risk_register__register_name').annotate(count=Count('id')),
    }
    html_string = render_to_string('reports/risk_register_summary.html', {
        'organization': org,
        'summary': summary,
        'filters': {'category': category, 'owner': owner, 'status': status, 'register': register},
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_risk_register_summary.pdf"'
    return response

def risk_heatmap_pdf(request):
    org = request.tenant
    risks = Risk.objects.filter(organization=org)
    
    # Apply register filter if provided
    register_filter = request.GET.get('register')
    if register_filter:
        risks = risks.filter(risk_register__register_name__icontains=register_filter)
    
    # Enhanced data preparation for bubble chart
    impact_range = range(1, 6)
    likelihood_range = range(1, 6)
    
    # Create bubble chart data with risk counts and details
    bubble_data = []
    risk_details = {}
    
    for impact in impact_range:
        for likelihood in likelihood_range:
            risk_count = risks.filter(
                residual_impact_score=impact, 
                residual_likelihood_score=likelihood
            ).count()
            
            if risk_count > 0:
                # Get actual risks for this cell
                cell_risks = risks.filter(
                    residual_impact_score=impact, 
                    residual_likelihood_score=likelihood
                )
                
                # Calculate risk score (impact * likelihood)
                risk_score = impact * likelihood
                
                bubble_data.append({
                    'impact': impact,
                    'likelihood': likelihood,
                    'count': risk_count,
                    'risk_score': risk_score,
                    'risks': list(cell_risks.values('risk_name', 'category', 'status', 'risk_owner')[:5])  # Top 5 risks
                })
                
                risk_details[f"{impact}-{likelihood}"] = {
                    'count': risk_count,
                    'risks': list(cell_risks.values('risk_name', 'category', 'status', 'risk_owner')),
                    'risk_score': risk_score
                }
    
    # Calculate comprehensive statistics using model risk levels to avoid gaps
    total_risks = risks.count()
    low_risk_count = 0
    medium_risk_count = 0
    high_risk_count = 0
    for r in risks:
        level = r.get_risk_level()
        if level == 'low':
            low_risk_count += 1
        elif level == 'medium':
            medium_risk_count += 1
        elif level in ('high', 'very_high', 'critical'):
            high_risk_count += 1
        else:
            # Fallback if level is None: infer with conservative thresholds
            score = r.residual_risk_score
            if score <= 5:
                low_risk_count += 1
            elif score <= 10:
                medium_risk_count += 1
            else:
                high_risk_count += 1
    
    # Risk distribution by category
    category_distribution = risks.values('category').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # Top risk areas (highest risk scores)
    top_risk_areas = []
    for data in bubble_data:
        if data['risk_score'] >= 12:  # High risk threshold
            top_risk_areas.append({
                'impact': data['impact'],
                'likelihood': data['likelihood'],
                'count': data['count'],
                'risk_score': data['risk_score'],
                'risks': data['risks']
            })
    
    # Generate enhanced bubble chart
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 8))
    
    # Main bubble chart
    if bubble_data:
        x = [d['impact'] for d in bubble_data]
        y = [d['likelihood'] for d in bubble_data]
        sizes = [d['count'] * 100 for d in bubble_data]  # Scale bubble sizes
        colors = [d['risk_score'] for d in bubble_data]
        
        scatter = ax1.scatter(x, y, s=sizes, c=colors, cmap='RdYlGn_r', alpha=0.7, edgecolors='black', linewidth=1)
        
        # Add risk count labels on bubbles
        for i, data in enumerate(bubble_data):
            if data['count'] > 0:
                ax1.annotate(str(data['count']), 
                           (data['impact'], data['likelihood']), 
                           ha='center', va='center', fontweight='bold', fontsize=10)
    
    ax1.set_xlabel('Impact Score', fontsize=12, fontweight='bold')
    ax1.set_ylabel('Likelihood Score', fontsize=12, fontweight='bold')
    ax1.set_title('Risk Distribution Bubble Chart', fontsize=14, fontweight='bold')
    ax1.set_xlim(0.5, 5.5)
    ax1.set_ylim(0.5, 5.5)
    ax1.grid(True, alpha=0.3)
    ax1.set_xticks(range(1, 6))
    ax1.set_yticks(range(1, 6))
    
    # Add colorbar
    cbar = plt.colorbar(scatter, ax=ax1)
    cbar.set_label('Risk Score (Impact × Likelihood)', fontsize=10)
    
    # Risk level zones
    ax1.axhline(y=3.5, color='red', linestyle='--', alpha=0.5, label='High Risk Zone')
    ax1.axvline(x=3.5, color='red', linestyle='--', alpha=0.5)
    ax1.axhline(y=2.5, color='orange', linestyle='--', alpha=0.5, label='Medium Risk Zone')
    ax1.axvline(x=2.5, color='orange', linestyle='--', alpha=0.5)
    ax1.legend()
    
    # Risk distribution pie chart
    if total_risks > 0:
        labels = ['High Risk', 'Medium Risk', 'Low Risk']
        sizes = [high_risk_count, medium_risk_count, low_risk_count]
        colors = ['#ff4444', '#ffaa00', '#44aa44']
        
        wedges, texts, autotexts = ax2.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', 
                                          startangle=90, explode=(0.1, 0.05, 0.05))
        ax2.set_title('Risk Level Distribution', fontsize=14, fontweight='bold')
        
        # Enhance text appearance
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
    
    plt.tight_layout()
    
    # Save chart
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    image_base64 = base64.b64encode(buf.read()).decode('utf-8')
    
    # Generate comprehensive context
    context = {
        'organization': org,
        'image_base64': image_base64,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Enterprise Risk Heatmap Analysis',
        'description': 'Comprehensive risk distribution analysis with bubble chart visualization',
        'total_risks': total_risks,
        'high_risk_count': high_risk_count,
        'medium_risk_count': medium_risk_count,
        'low_risk_count': low_risk_count,
        'category_distribution': category_distribution,
        'top_risk_areas': top_risk_areas,
        'risk_details': risk_details,
        'bubble_data': bubble_data,
        'filters_summary': f"Register: {register_filter}" if register_filter else 'All risks analyzed'
    }
    
    html_string = render_to_string('reports/risk_heatmap.html', context)
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_enterprise_risk_heatmap.pdf"'
    return response

def risk_trends_pdf(request):
    org = request.tenant
    risks = Risk.objects.filter(organization=org)
    # Group by month using TruncMonth for DB compatibility
    risks_by_month = (
        risks.annotate(month=TruncMonth('date_identified'))
             .values('month')
             .annotate(count=Count('id'))
             .order_by('month')
    )
    months = [r['month'].strftime('%Y-%m') if r['month'] else '' for r in risks_by_month]
    counts = [r['count'] for r in risks_by_month]
    # Generate trend chart
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(months, counts, marker='o')
    ax.set_xlabel('Month')
    ax.set_ylabel('Number of Risks Identified')
    ax.set_title('Risk Identification Trend')
    plt.xticks(rotation=45)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)
    image_base64 = base64.b64encode(buf.read()).decode('utf-8')
    html_string = render_to_string('reports/risk_trends.html', {
        'organization': org,
        'image_base64': image_base64,
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_risk_trends.pdf"'
    return response

def control_effectiveness_pdf(request):
    org = request.tenant
    controls = Control.objects.filter(organization=org)
    effectiveness = controls.values('effectiveness_rating').annotate(count=Count('id'))
    labels = [e['effectiveness_rating'] for e in effectiveness]
    counts = [e['count'] for e in effectiveness]
    # Pie chart
    fig, ax = plt.subplots()
    ax.pie(counts, labels=labels, autopct='%1.1f%%', startangle=90)
    ax.set_title('Control Effectiveness Distribution')
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)
    image_base64 = base64.b64encode(buf.read()).decode('utf-8')
    html_string = render_to_string('reports/control_effectiveness.html', {
        'organization': org,
        'image_base64': image_base64,
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_control_effectiveness.pdf"'
    return response

def kri_status_pdf(request):
    org = request.tenant
    kris = KRI.objects.filter(risk__organization=org)
    status_counts = {'normal': 0, 'warning': 0, 'critical': 0}
    for kri in kris:
        status = kri.get_status()
        if status in status_counts:
            status_counts[status] += 1
    html_string = render_to_string('reports/kri_status.html', {
        'organization': org,
        'status_counts': status_counts,
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_kri_status.pdf"'
    return response

def assessment_timeline_pdf(request):
    org = request.tenant
    assessments = RiskAssessment.objects.filter(organization=org)
    # Group by date
    timeline = assessments.values('assessment_date').annotate(avg_score=Count('risk_score')).order_by('assessment_date')
    dates = [t['assessment_date'] for t in timeline]
    scores = [t['avg_score'] for t in timeline]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(dates, scores, marker='o')
    ax.set_xlabel('Date')
    ax.set_ylabel('Avg Risk Score')
    ax.set_title('Assessment Timeline')
    plt.xticks(rotation=45)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)
    image_base64 = base64.b64encode(buf.read()).decode('utf-8')
    html_string = render_to_string('reports/assessment_timeline.html', {
        'organization': org,
        'image_base64': image_base64,
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_assessment_timeline.pdf"'
    return response

def risk_register_detailed_pdf(request):
    org = request.tenant
    risks = Risk.objects.filter(organization=org)
    category = request.GET.get('category')
    owner = request.GET.get('owner')
    status = request.GET.get('status')
    register = request.GET.get('register')
    if category:
        risks = risks.filter(category=category)
    if owner:
        risks = risks.filter(risk_owner__icontains=owner)
    if status:
        risks = risks.filter(status=status)
    if register:
        risks = risks.filter(risk_register__register_name__icontains=register)
    
    # Calculate additional statistics
    total_risks = risks.count()
    high_priority_risks = risks.filter(residual_risk_score__gte=15).count()
    active_risks = risks.exclude(status='closed').exclude(status='archived').count()
    categories_count = risks.values('category').distinct().count()
    
    html_string = render_to_string('reports/risk_register_detailed.html', {
        'organization': org,
        'risks': risks,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Risk Register - Detailed Report',
        'description': f'Comprehensive risk analysis with {total_risks} risks identified',
        'filters': {
            'category': category,
            'owner': owner,
            'status': status,
            'register': register
        },
        'filters_summary': ', '.join([f"{k}: {v}" for k, v in {'category': category, 'owner': owner, 'status': status, 'register': register}.items() if v]),
        'total_risks': total_risks,
        'high_priority_risks': high_priority_risks,
        'active_risks': active_risks,
        'categories_count': categories_count,
        'for_pdf': True
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_risk_register_detailed.pdf"'
    return response

def risk_assessment_details_pdf(request):
    org = request.tenant
    assessments = RiskAssessment.objects.filter(organization=org)
    risk_id = request.GET.get('risk')
    assessment_type = request.GET.get('assessment_type')
    assessor = request.GET.get('assessor')
    min_score = request.GET.get('min_score')
    max_score = request.GET.get('max_score')
    if risk_id:
        assessments = assessments.filter(risk_id=risk_id)
    if assessment_type:
        assessments = assessments.filter(assessment_type=assessment_type)
    if assessor:
        assessments = assessments.filter(assessor__icontains=assessor)
    if min_score:
        assessments = assessments.filter(risk_score__gte=min_score)
    if max_score:
        assessments = assessments.filter(risk_score__lte=max_score)
    
    # Calculate additional statistics
    total_assessments = assessments.count()
    high_risk_assessments = assessments.filter(risk_score__gte=15).count()
    medium_risk_assessments = assessments.filter(risk_score__range=(5, 14)).count()
    assessors = assessments.values('assessor').distinct().count()
    
    html_string = render_to_string('reports/risk_assessment_details.html', {
        'organization': org,
        'assessments': assessments,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Risk Assessment Analysis',
        'description': f'Risk assessment analysis with {total_assessments} assessments',
        'filters': {
            'risk': risk_id,
            'assessment_type': assessment_type,
            'assessor': assessor,
            'min_score': min_score,
            'max_score': max_score
        },
        'filters_summary': ', '.join([f"{k}: {v}" for k, v in {'risk': risk_id, 'assessment_type': assessment_type, 'assessor': assessor, 'min_score': min_score, 'max_score': max_score}.items() if v]),
        'total_assessments': total_assessments,
        'high_risk_assessments': high_risk_assessments,
        'medium_risk_assessments': medium_risk_assessments,
        'assessors': assessors,
        'for_pdf': True
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_risk_assessment_details.pdf"'
    return response

def control_details_pdf(request):
    org = request.tenant
    controls = Control.objects.filter(organization=org)
    control_type_filter = request.GET.get('control_type')
    status_filter = request.GET.get('status')
    owner_filter = request.GET.get('owner')
    
    if control_type_filter:
        controls = controls.filter(control_type=control_type_filter)
    if status_filter:
        controls = controls.filter(status=status_filter)
    if owner_filter:
        controls = controls.filter(control_owner__icontains=owner_filter)
    
    # Calculate additional statistics
    total_controls = controls.count()
    effective_controls = controls.filter(effectiveness_rating='effective').count()
    ineffective_controls = controls.filter(effectiveness_rating='ineffective').count()
    control_owners = controls.values('control_owner').distinct().count()
    
    html_string = render_to_string('reports/control_details.html', {
        'organization': org,
        'controls': controls,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Control Effectiveness Analysis',
        'description': f'Control effectiveness analysis with {total_controls} controls',
        'filters': {
            'control_type': control_type_filter,
            'status': status_filter,
            'owner': owner_filter
        },
        'filters_summary': ', '.join([f"{k}: {v}" for k, v in {'control_type': control_type_filter, 'status': status_filter, 'owner': owner_filter}.items() if v]),
        'total_controls': total_controls,
        'effective_controls': effective_controls,
        'ineffective_controls': ineffective_controls,
        'control_owners': control_owners,
        'for_pdf': True
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_control_details.pdf"'
    return response

def kri_details_pdf(request):
    org = request.tenant
    kris = KRI.objects.filter(risk__organization=org)
    risk_filter = request.GET.get('risk')
    status_filter = request.GET.get('status')
    direction_filter = request.GET.get('direction')
    
    if risk_filter:
        kris = kris.filter(risk__risk_name__icontains=risk_filter)
    if status_filter:
        kris = kris.filter(status=status_filter)
    if direction_filter:
        kris = kris.filter(direction=direction_filter)
    
    # Calculate additional statistics
    total_kris = kris.count()
    critical_kris = sum(1 for kri in kris if kri.get_status() == 'critical')
    warning_kris = sum(1 for kri in kris if kri.get_status() == 'warning')
    monitored_risks = kris.values('risk').distinct().count()
    
    html_string = render_to_string('reports/kri_details.html', {
        'organization': org,
        'kris': kris,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Key Risk Indicators Analysis',
        'description': f'KRI monitoring analysis with {total_kris} indicators',
        'filters': {
            'risk': risk_filter,
            'status': status_filter,
            'direction': direction_filter
        },
        'filters_summary': ', '.join([f"{k}: {v}" for k, v in {'risk': risk_filter, 'status': status_filter, 'direction': direction_filter}.items() if v]),
        'total_kris': total_kris,
        'critical_kris': critical_kris,
        'warning_kris': warning_kris,
        'monitored_risks': monitored_risks,
        'for_pdf': True
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_kri_details.pdf"'
    return response

def workplan_summary_pdf(request):
    org = request.tenant
    year = request.GET.get('year')
    status = request.GET.get('status')
    workplans = AuditWorkplan.objects.filter(organization=org)
    if year:
        workplans = workplans.filter(fiscal_year=year)
    if status:
        workplans = workplans.filter(state=status)
    summary = workplans.values('fiscal_year', 'state').annotate(count=Count('id')).order_by('-fiscal_year')
    
    # Calculate additional statistics
    total_workplans = workplans.count()
    active_workplans = workplans.exclude(state='cancelled').count()
    fiscal_years = workplans.values('fiscal_year').distinct().count()
    
    context = {
        'organization': org,
        'summary': summary,
        'filters': {'year': year, 'status': status},
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Audit Workplan Summary Report',
        'description': f'Comprehensive workplan analysis with {total_workplans} workplans across {fiscal_years} fiscal years',
        'filters_summary': ', '.join([f"{k}: {v}" for k, v in {'year': year, 'status': status}.items() if v]),
        'total_workplans': total_workplans,
        'active_workplans': active_workplans,
        'fiscal_years': fiscal_years,
    }

    if request.GET.get('format') == 'docx':
        gen_ts = context['generation_timestamp']
        doc = _docx_start_document(org, context['title'], gen_ts)
        _docx_add_heading(doc, 'Summary by Fiscal Year and Status')
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Fiscal Year'
        hdr[1].text = 'Status'
        hdr[2].text = 'Count'
        for row in context['summary']:
            r = tbl.add_row().cells
            r[0].text = str(row.get('fiscal_year'))
            r[1].text = str(row.get('state'))
            r[2].text = str(row.get('count'))
        return _docx_http_response(doc, 'audit_workplan_summary', org)
    else:
        html_string = render_to_string('reports/audit_workplan_summary.html', context)
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_audit_workplan_summary.pdf"'
    return response

def get_engagement_names(org):
    return list(Engagement.objects.filter(organization=org).values_list('title', flat=True).distinct().order_by('title'))

def engagement_summary_pdf(request):
    org = request.tenant
    status = request.GET.get('status')
    assigned_to = request.GET.get('assigned_to')
    engagement_name = request.GET.get('engagement_name')
    engagements = Engagement.objects.filter(organization=org)
    if status:
        engagements = engagements.filter(project_status=status)
    if assigned_to:
        engagements = engagements.filter(assigned_to__email__icontains=assigned_to)
    if engagement_name:
        # Try exact match first, then icontains
        if engagements.filter(title=engagement_name).exists():
            engagements = engagements.filter(title=engagement_name)
        else:
            engagements = engagements.filter(title__icontains=engagement_name)
    summary = engagements.values('project_status', 'assigned_to__email').annotate(count=Count('id')).order_by('project_status')
    engagement_names = get_engagement_names(org)
    
    # Calculate additional statistics
    total_engagements = engagements.count()
    active_engagements = engagements.exclude(project_status__in=['completed', 'cancelled']).count()
    assigned_auditors = engagements.values('assigned_to__email').distinct().count()
    
    context = {
        'organization': org,
        'summary': summary,
        'filters': {'status': status, 'assigned_to': assigned_to, 'engagement_name': engagement_name},
        'engagement_names': engagement_names,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Audit Engagement Summary Report',
        'description': f'Comprehensive engagement analysis with {total_engagements} engagements and {assigned_auditors} assigned auditors',
        'filters_summary': ', '.join([f"{k}: {v}" for k, v in {'status': status, 'assigned_to': assigned_to, 'engagement_name': engagement_name}.items() if v]),
        'total_engagements': total_engagements,
        'active_engagements': active_engagements,
        'assigned_auditors': assigned_auditors,
    }
    if request.GET.get('format') == 'docx':
        gen_ts = context['generation_timestamp']
        doc = _docx_start_document(org, context['title'], gen_ts)
        _docx_add_heading(doc, 'Engagement Summary')
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Status'
        hdr[1].text = 'Assigned To (email)'
        hdr[2].text = 'Count'
        for row in context['summary']:
            r = tbl.add_row().cells
            r[0].text = str(row.get('project_status'))
            r[1].text = str(row.get('assigned_to__email'))
            r[2].text = str(row.get('count'))
        return _docx_http_response(doc, 'audit_engagement_summary', org)
    else:
        html_string = render_to_string('reports/audit_engagement_summary.html', context)
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_audit_engagement_summary.pdf"'
    return response

def issue_register_pdf(request):
    org = request.tenant
    status = request.GET.get('status')
    severity = request.GET.get('severity')  # risk_level param
    engagement_name = request.GET.get('engagement_name')
    issues = Issue.objects.filter(organization=org)
    if status:
        issues = issues.filter(issue_status=status)
    if severity:
        issues = issues.filter(risk_level=severity)
    if engagement_name:
        # Filter issues by engagement through the relationship chain: Issue -> Procedure -> Risk -> Objective -> Engagement
        if Engagement.objects.filter(organization=org, title=engagement_name).exists():
            issues = issues.filter(procedure__risk__objective__engagement__title=engagement_name)
        else:
            issues = issues.filter(procedure__risk__objective__engagement__title__icontains=engagement_name)
    engagement_names = get_engagement_names(org)
    
    # Calculate additional statistics
    total_issues = issues.count()
    open_issues = issues.exclude(issue_status='closed').count()
    high_priority_issues = issues.filter(risk_level__in=['high', 'critical']).count()
    escalated_issues = issues.filter(issue_status='escalated').count()
    
    context = {
        'organization': org,
        'issues': issues,
        'filters': {'status': status, 'severity': severity, 'engagement_name': engagement_name},
        'engagement_names': engagement_names,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Audit Issue Register Report',
        'description': f'Comprehensive audit issues analysis with {total_issues} issues identified',
        'filters_summary': ', '.join([f"{k}: {v}" for k, v in {'status': status, 'severity': severity, 'engagement_name': engagement_name}.items() if v]),
        'total_issues': total_issues,
        'open_issues': open_issues,
        'high_priority_issues': high_priority_issues,
        'escalated_issues': escalated_issues,
    }
    if request.GET.get('format') == 'docx':
        gen_ts = context['generation_timestamp']
        doc = _docx_start_document(org, context['title'], gen_ts)
        _docx_add_heading(doc, 'Issues')
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Issue ID'
        hdr[1].text = 'Title'
        hdr[2].text = 'Status'
        hdr[3].text = 'Risk Level'
        hdr[4].text = 'Owner'
        hdr[5].text = 'Target Date'
        for issue in context['issues']:
            r = tbl.add_row().cells
            r[0].text = str(getattr(issue, 'issue_id', getattr(issue, 'id', '')))
            r[1].text = str(getattr(issue, 'title', ''))
            r[2].text = str(getattr(issue, 'issue_status', ''))
            r[3].text = str(getattr(issue, 'risk_level', ''))
            r[4].text = str(getattr(getattr(issue, 'issue_owner', None), 'get_full_name', lambda: '')() or getattr(issue, 'issue_owner_title', '') )
            r[5].text = str(getattr(issue, 'target_date', ''))
        return _docx_http_response(doc, 'audit_issue_register', org)
    else:
        html_string = render_to_string('reports/audit_issue_register.html', context)
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_audit_issue_register.pdf"'
    return response

def issue_followup_pdf(request):
    org = request.tenant
    overdue = request.GET.get('overdue')
    issues = Issue.objects.filter(organization=org)
    
    # Apply overdue filter if specified
    if overdue == '1':
        today = timezone.now().date()
        issues = issues.filter(target_date__lt=today, issue_status__in=['open', 'in_progress'])
    
    # Calculate comprehensive statistics
    total_issues = issues.count()
    open_issues = issues.filter(issue_status='open').count()
    in_progress_issues = issues.filter(issue_status='in_progress').count()
    closed_issues = issues.filter(issue_status='closed').count()
    
    # Risk level distribution
    high_risk_issues = issues.filter(risk_level='high').count()
    medium_risk_issues = issues.filter(risk_level='medium').count()
    low_risk_issues = issues.filter(risk_level='low').count()
    critical_risk_issues = issues.filter(risk_level='critical').count()
    
    # Overdue analysis
    overdue_issues = issues.filter(
        target_date__lt=timezone.now().date(),
        issue_status__in=['open', 'in_progress']
    ).count()
    
    # Issue type distribution
    issue_type_distribution = issues.values('issue_type').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # Top issue owners
    top_owners = issues.values('issue_owner__first_name', 'issue_owner__last_name', 'issue_owner_email').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # Average days overdue
    overdue_issues_data = issues.filter(
        target_date__lt=timezone.now().date(),
        issue_status__in=['open', 'in_progress']
    )
    avg_days_overdue = overdue_issues_data.aggregate(
        avg_days=Avg('days_overdue')
    )['avg_days'] or 0
    
    # Financial impact analysis
    total_financial_impact = issues.aggregate(
        total_impact=Sum('financial_impact')
    )['total_impact'] or 0
    
    # Generate comprehensive context
    context = {
        'organization': org,
        'issues': issues,
        'filters': {'overdue': overdue},
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Audit Issue Follow-up Report',
        'description': 'Comprehensive audit issue tracking and follow-up analysis',
        'total_issues': total_issues,
        'open_issues': open_issues,
        'in_progress_issues': in_progress_issues,
        'closed_issues': closed_issues,
        'high_risk_issues': high_risk_issues,
        'medium_risk_issues': medium_risk_issues,
        'low_risk_issues': low_risk_issues,
        'critical_risk_issues': critical_risk_issues,
        'overdue_issues': overdue_issues,
        'issue_type_distribution': issue_type_distribution,
        'top_owners': top_owners,
        'avg_days_overdue': round(avg_days_overdue, 1),
        'total_financial_impact': total_financial_impact,
        'filters_summary': f"Overdue Only: {'Yes' if overdue == '1' else 'No'}"
    }
    
    if request.GET.get('format') == 'docx':
        doc = _docx_start_document(org, context['title'], context['generation_timestamp'])
        _docx_add_heading(doc, 'Issues Pending Follow-up')
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Issue ID'
        hdr[1].text = 'Title'
        hdr[2].text = 'Status'
        hdr[3].text = 'Owner'
        hdr[4].text = 'Target Date'
        for issue in issues:
            r = tbl.add_row().cells
            r[0].text = str(getattr(issue, 'code', getattr(issue, 'id', '')))
            r[1].text = str(getattr(issue, 'issue_title', ''))
            r[2].text = str(getattr(issue, 'issue_status', ''))
            r[3].text = str(getattr(getattr(issue, 'issue_owner', None), 'get_full_name', lambda: '')() or getattr(issue, 'issue_owner_title', '') )
            r[4].text = str(getattr(issue, 'target_date', ''))
        return _docx_http_response(doc, 'audit_issue_followup', org)
    else:
        html_string = render_to_string('reports/audit_issue_followup.html', context)
        pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
            stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
        )
        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{org.code}_audit_issue_followup.pdf"'
        return response

def approval_workflow_pdf(request):
    org = request.tenant
    status = request.GET.get('status')
    approvals = Approval.objects.filter(organization=org)
    if status:
        approvals = approvals.filter(status=status)
    context = {
        'organization': org,
        'approvals': approvals,
        'filters': {'status': status},
    }
    if request.GET.get('format') == 'docx':
        gen_ts = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        doc = _docx_start_document(org, 'Audit Approval Workflow Report', gen_ts)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Approval ID'
        hdr[1].text = 'Status'
        hdr[2].text = 'Created'
        hdr[3].text = 'Updated'
        for a in approvals:
            r = tbl.add_row().cells
            r[0].text = str(getattr(a, 'id', ''))
            r[1].text = str(getattr(a, 'status', ''))
            r[2].text = str(getattr(a, 'created_at', ''))
            r[3].text = str(getattr(a, 'updated_at', ''))
        return _docx_http_response(doc, 'audit_approval_workflow', org)
    else:
        html_string = render_to_string('reports/audit_approval_workflow.html', context)
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_audit_approval_workflow.pdf"'
    return response

def smart_engagement_progress_pdf(request):
    org = request.tenant
    # Enhanced engagement analysis with comprehensive statistics
    from django.db.models import Avg, Count, F, ExpressionWrapper, DurationField
    from django.db.models.functions import Now
    
    engagements = Engagement.objects.filter(organization=org)
    
    # Basic statistics
    total_engagements = engagements.count()
    closed_engagements = engagements.filter(project_status='closed').count()
    in_progress_engagements = engagements.filter(project_status='in_progress').count()
    planned_engagements = engagements.filter(project_status='planned').count()
    
    # Calculate percentages
    percent_closed = (closed_engagements / total_engagements * 100) if total_engagements > 0 else 0
    
    # Duration analysis
    completed_engagements = engagements.filter(
        project_status='closed',
        target_end_date__isnull=False,
        project_start_date__isnull=False
    )
    
    avg_duration = completed_engagements.annotate(
        duration=ExpressionWrapper(F('target_end_date') - F('project_start_date'), output_field=DurationField())
    ).aggregate(avg=Avg('duration'))['avg']
    
    # Convert timedelta to days for display
    avg_duration_days = avg_duration.days if avg_duration else 0
    
    # Engagement type distribution
    engagement_type_distribution = engagements.values('engagement_type').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # Top assigned auditors
    top_auditors = engagements.values('assigned_to__first_name', 'assigned_to__last_name').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # Workplan distribution
    workplan_distribution = engagements.values('annual_workplan__name').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # Timeline analysis
    overdue_engagements = engagements.filter(
        target_end_date__lt=timezone.now().date(),
        project_status__in=['in_progress', 'planned']
    ).count()
    
    # Risk assessment analysis based on conclusion
    high_risk_engagements = engagements.filter(conclusion__in=['unsatisfactory', 'significant_improvement_needed']).count()
    medium_risk_engagements = engagements.filter(conclusion='needs_improvement').count()
    low_risk_engagements = engagements.filter(conclusion='satisfactory').count()
    
    # Generate comprehensive context
    context = {
        'organization': org,
        'total_engagements': total_engagements,
        'closed_engagements': closed_engagements,
        'in_progress_engagements': in_progress_engagements,
        'planned_engagements': planned_engagements,
        'percent_closed': round(percent_closed, 2),
        'avg_duration_days': avg_duration_days,
        'engagement_type_distribution': engagement_type_distribution,
        'top_auditors': top_auditors,
        'workplan_distribution': workplan_distribution,
        'overdue_engagements': overdue_engagements,
        'high_risk_engagements': high_risk_engagements,
        'medium_risk_engagements': medium_risk_engagements,
        'low_risk_engagements': low_risk_engagements,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Smart Engagement Progress Report',
        'description': 'Comprehensive audit engagement progress analysis and performance metrics',
        'filters_summary': 'All engagements analyzed'
    }
    
    if request.GET.get('format') == 'docx':
        gen_ts = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        doc = _docx_start_document(org, 'Audit Engagement Progress Report', gen_ts)
        _docx_add_heading(doc, 'Progress Summary')
        tbl = doc.add_table(rows=3, cols=2)
        tbl.style = 'Table Grid'
        tbl.cell(0,0).text = 'Total engagements'; tbl.cell(0,1).text = str(total_engagements)
        tbl.cell(1,0).text = 'Closed engagements'; tbl.cell(1,1).text = str(closed_engagements)
        tbl.cell(2,0).text = 'Percent closed'; tbl.cell(2,1).text = f"{percent_closed:.2f}%"
        return _docx_http_response(doc, 'audit_engagement_progress', org)
    else:
        html_string = render_to_string('reports/smart_engagement_progress.html', context)
        pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
            stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
        )
        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{org.code}_smart_engagement_progress.pdf"'
        return response

def get_engagement_by_name(org, engagement_name):
    if engagement_name:
        # Try exact match first
        engagement = Engagement.objects.filter(organization=org, title=engagement_name).first()
        if not engagement:
            # Fallback to icontains for partial/typed input
            engagement = Engagement.objects.filter(organization=org, title__icontains=engagement_name).first()
        return engagement
    return None

def engagement_details_pdf(request):
    org = request.tenant
    engagement_name = request.GET.get('engagement_name') or request.GET.get('q')
    engagement = None
    if engagement_name:
        # Try exact match first, then icontains
        engagement = Engagement.objects.filter(organization=org, title=engagement_name).first()
        if not engagement:
            engagement = Engagement.objects.filter(organization=org, title__icontains=engagement_name).first()
    else:
        # If no filter, show the most recent engagement
        engagement = Engagement.objects.filter(organization=org).order_by('-project_start_date').first()

    engagement_names = get_engagement_names(org)

    # Annex 1: Priorities
    annex_recommendation_priorities = list(getattr(Recommendation, 'PRIORITY_CHOICES', []))
    annex_priority_definitions = {
        'low': 'Action is desirable; failure to act is unlikely to cause material negative consequences.',
        'medium': 'Action is required to ensure the entity is not exposed to risks; failure to act could result in negative consequences.',
        'high': 'Prompt action is required to ensure the entity is not exposed to high risks; failure to act could result in major negative consequences.',
        'critical': 'Immediate action is imperative; risk exposure is critical and may lead to severe consequences if not addressed.',
    }
    priority_definitions = [
        {
            "priority": val,
            "label": label,
            "definition": annex_priority_definitions.get(val, "Priority definition"),
        }
        for val, label in annex_recommendation_priorities
    ]

    # Annex 2: Conclusion ratings
    annex_conclusion_choices = list(getattr(Engagement, 'CONCLUSION_CHOICES', []))
    annex_conclusion_definitions = {
        'satisfactory': 'Governance, risk management and control processes are adequately designed and operating effectively.',
        'needs_improvement': 'Arrangements are generally established, but improvements are needed to address noted gaps that do not significantly impair objectives.',
        'significant_improvement_needed': 'Key weaknesses exist; major improvement is needed and results may materially affect achievement of objectives.',
        'unsatisfactory': 'Governance arrangements and controls are not adequately established or not functioning, posing significant risk to objectives.',
        'not_rated': 'An overall rating was not assigned for this engagement.',
    }
    conclusion_definitions = [
        {
            "conclusion": val,
            "label": label,
            "definition": annex_conclusion_definitions.get(val, "Rating definition"),
        }
        for val, label in annex_conclusion_choices
    ]

    # Annex 3: Risk levels (same approach if needed later)

    context = {
        'organization': org,
        'engagement': engagement,
        'engagement_names': engagement_names,
        'filters': {'engagement_name': engagement_name},
        'for_pdf': True,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Audit Engagement Details Report',
        'description': f'Detailed analysis of engagement: {engagement.title if engagement else "Not found"}',

        # Clean lists for template
        'priority_definitions': priority_definitions,
        'conclusion_definitions': conclusion_definitions,

        # Still keep originals in case other parts use them
        'annex_recommendation_priorities': annex_recommendation_priorities,
        'annex_priority_definitions': annex_priority_definitions,
        'annex_conclusion_choices': annex_conclusion_choices,
        'annex_conclusion_definitions': annex_conclusion_definitions,
        'annex_issue_risk_levels': list(getattr(Issue, 'RISK_LEVEL_CHOICES', [])),
    }

    if request.GET.get('format') == 'docx':
        ctx = context
        doc = _docx_start_document(org, ctx['title'], ctx['generation_timestamp'])
        if engagement:
            # Basic Information table (mirrors PDF)
            tbl = doc.add_table(rows=5, cols=2)
            tbl.style = 'Table Grid'
            tbl.cell(0,0).text = 'Engagement Title'; tbl.cell(0,1).text = engagement.title
            tbl.cell(1,0).text = 'Status'; tbl.cell(1,1).text = getattr(engagement, 'get_project_status_display', lambda: getattr(engagement, 'project_status',''))()
            tbl.cell(2,0).text = 'Workplan'; tbl.cell(2,1).text = (getattr(getattr(engagement,'annual_workplan',None),'name',None) or 'Not specified')
            tbl.cell(3,0).text = 'Start Date'; tbl.cell(3,1).text = str(getattr(engagement, 'project_start_date', ''))
            tbl.cell(4,0).text = 'Report Issued Date'; tbl.cell(4,1).text = str(getattr(engagement, 'report_issued_date', ''))

            # Executive Summary
            if getattr(engagement, 'executive_summary', None):
                _docx_add_heading(doc, 'Executive Summary')
                _docx_add_html_block(doc, getattr(engagement, 'executive_summary', ''))

            # Purpose and Background
            if getattr(engagement, 'purpose', None):
                _docx_add_heading(doc, 'Purpose')
                _docx_add_html_block(doc, getattr(engagement, 'purpose', ''))
            if getattr(engagement, 'background', None):
                _docx_add_heading(doc, 'Background')
                _docx_add_html_block(doc, getattr(engagement, 'background', ''))

            # Objectives list
            objectives = getattr(engagement, 'objectives', None)
            if objectives and hasattr(objectives, 'all'):
                _docx_add_heading(doc, 'Audit Objectives')
                for obj in objectives.all():
                    doc.add_paragraph(f"- {getattr(obj,'title','')}")

            # Conclusion
            _docx_add_heading(doc, 'Conclusion')
            conclusion_text = getattr(engagement, 'get_conclusion_display', lambda: getattr(engagement, 'conclusion',''))()
            doc.add_paragraph(str(conclusion_text))
            if getattr(engagement, 'conclusion_description', None):
                _docx_add_html_block(doc, getattr(engagement, 'conclusion_description', ''))

            # Issues and Findings
            issues = getattr(engagement, 'all_issues', []) or []
            if issues:
                _docx_add_heading(doc, 'Issues and Findings')
                for idx, issue in enumerate(issues, start=1):
                    doc.add_paragraph(f"Issue {idx}: {getattr(issue,'issue_title','')}")
                    if getattr(issue, 'issue_description', None):
                        _docx_add_html_block(doc, f"Description: {getattr(issue,'issue_description','')}")
                    if getattr(issue, 'root_cause', None):
                        _docx_add_html_block(doc, f"Root Cause: {getattr(issue,'root_cause','')}")
                    recs = getattr(issue, 'recommendations', None)
                    if recs and hasattr(recs, 'all'):
                        _docx_add_heading(doc, 'Recommendations:')
                        for rec in recs.all():
                            doc.add_paragraph(f"• {getattr(rec,'title','')}")
                            if getattr(rec, 'description', None):
                                _docx_add_html_block(doc, f"{getattr(rec,'description','')}")
                    if getattr(issue, 'risks', None):
                        _docx_add_html_block(doc, f"Risks: {getattr(issue,'risks','')}")
                    if getattr(issue, 'management_action_plan', None):
                        _docx_add_html_block(doc, f"Management Action Plan: {getattr(issue,'management_action_plan','')}")
        else:
            doc.add_paragraph('No engagement found for the provided filter.')
        return _docx_http_response(doc, 'audit_engagement_details', org)
    else:
        html_string = render_to_string('reports/audit_engagement_details.html', context)
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_audit_engagement_details.pdf"'
    return response

def engagement_with_issues_pdf(request):
    org = request.tenant
    engagement_name = request.GET.get('engagement_name') or request.GET.get('q')
    engagement = None
    
    if engagement_name:
        engagement = Engagement.objects.filter(organization=org, title=engagement_name).first()
        if not engagement:
            engagement = Engagement.objects.filter(organization=org, title__icontains=engagement_name).first()
    else:
        engagement = Engagement.objects.filter(organization=org).order_by('-project_start_date').first()
    
    issues = engagement.all_issues if engagement else []
    engagement_names = get_engagement_names(org)
    
    # Calculate comprehensive statistics
    total_issues = len(issues) if issues else 0
    open_issues = len([i for i in issues if i.issue_status == 'open']) if issues else 0
    closed_issues = len([i for i in issues if i.issue_status == 'closed']) if issues else 0
    in_progress_issues = len([i for i in issues if i.issue_status == 'in_progress']) if issues else 0
    
    # Risk level distribution
    high_risk_issues = len([i for i in issues if i.risk_level == 'high']) if issues else 0
    medium_risk_issues = len([i for i in issues if i.risk_level == 'medium']) if issues else 0
    low_risk_issues = len([i for i in issues if i.risk_level == 'low']) if issues else 0
    critical_risk_issues = len([i for i in issues if i.risk_level == 'critical']) if issues else 0
    
    # Issue type distribution
    issue_type_distribution = {}
    if issues:
        for issue in issues:
            issue_type = issue.issue_type or 'other'
            issue_type_distribution[issue_type] = issue_type_distribution.get(issue_type, 0) + 1
    
    # Overdue analysis
    overdue_issues = len([i for i in issues if i.target_date and i.target_date < timezone.now().date() and i.issue_status in ['open', 'in_progress']]) if issues else 0
    
    # Financial impact analysis
    total_financial_impact = sum([i.financial_impact or 0 for i in issues]) if issues else 0
    
    # Generate comprehensive context
    context = {
        'organization': org,
        'engagement': engagement,
        'issues': issues,
        'engagement_names': engagement_names,
        'filters': {'engagement_name': engagement_name},
        'for_pdf': True,  # Always set for PDF context
        'total_issues': total_issues,
        'open_issues': open_issues,
        'closed_issues': closed_issues,
        'in_progress_issues': in_progress_issues,
        'high_risk_issues': high_risk_issues,
        'medium_risk_issues': medium_risk_issues,
        'low_risk_issues': low_risk_issues,
        'critical_risk_issues': critical_risk_issues,
        'issue_type_distribution': issue_type_distribution,
        'overdue_issues': overdue_issues,
        'total_financial_impact': total_financial_impact,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Audit Engagement With Issues Report',
        'description': 'Comprehensive audit engagement analysis with detailed issue breakdown',
        'filters_summary': f"Engagement: {engagement.title if engagement else 'Not specified'}"
    }
    
    if request.GET.get('format') == 'docx':
        doc = _docx_start_document(org, 'Audit Engagement With Issues Report', timezone.now().strftime('%Y-%m-%d %H:%M:%S'))
        if engagement:
            _docx_add_heading(doc, f"Engagement: {engagement.title}")
        _docx_add_heading(doc, 'Issues')
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Issue ID'
        hdr[1].text = 'Title'
        hdr[2].text = 'Status'
        hdr[3].text = 'Risk Level'
        hdr[4].text = 'Target Date'
        for issue in issues:
            r = tbl.add_row().cells
            r[0].text = str(getattr(issue, 'code', getattr(issue, 'id', '')))
            r[1].text = str(getattr(issue, 'issue_title', ''))
            r[2].text = str(getattr(issue, 'issue_status', ''))
            r[3].text = str(getattr(issue, 'risk_level', ''))
            r[4].text = str(getattr(issue, 'target_date', ''))
        return _docx_http_response(doc, 'audit_engagement_with_issues', org)
    else:
        html_string = render_to_string('reports/audit_engagement_with_issues.html', context)
        pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
            stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
        )
        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{org.code}_audit_engagement_with_issues.pdf"'
        return response

def legal_case_summary_pdf(request):
    org = request.tenant
    status = request.GET.get('status')
    case_type = request.GET.get('case_type')
    priority = request.GET.get('priority')
    lead_attorney = request.GET.get('lead_attorney')
    case_name = request.GET.get('case_name')
    cases = LegalCase.objects.filter(organization=org)
    if status:
        cases = cases.filter(status=status)
    if case_type:
        cases = cases.filter(case_type__name__icontains=case_type)
    if priority:
        cases = cases.filter(priority=priority)
    if lead_attorney:
        cases = cases.filter(lead_attorney__email__icontains=lead_attorney)
    if case_name:
        cases = cases.filter(title__icontains=case_name)
    html_string = render_to_string('reports/legal_case_summary.html', {
        'organization': org,
        'cases': cases,
        'filters': {'status': status, 'case_type': case_type, 'priority': priority, 'lead_attorney': lead_attorney, 'case_name': case_name},
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_case_summary.pdf"'
    return response

def legal_case_details_pdf(request):
    org = request.tenant
    case_name = request.GET.get('case_name')
    case = None
    parties = []
    tasks = []
    documents = []
    if case_name:
        case = LegalCase.objects.filter(organization=org, title__icontains=case_name).first()
        if case:
            parties = case.parties.all()
            tasks = case.tasks.all()
            documents = case.documents.all()
    html_string = render_to_string('reports/legal_case_details.html', {
        'organization': org,
        'case': case,
        'parties': parties,
        'tasks': tasks,
        'documents': documents,
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_case_details.pdf"'
    return response

def compliance_requirement_summary_pdf(request):
    org = request.tenant
    requirements = ComplianceRequirement.objects.filter(organization=org)
    framework = request.GET.get('framework')
    jurisdiction = request.GET.get('jurisdiction')
    mandatory = request.GET.get('mandatory')
    policy = request.GET.get('policy')
    title = request.GET.get('title')
    if framework:
        requirements = requirements.filter(regulatory_framework__name__icontains=framework)
    if jurisdiction:
        requirements = requirements.filter(jurisdiction__icontains=jurisdiction)
    if mandatory in ['0', '1']:
        requirements = requirements.filter(mandatory=(mandatory == '1'))
    if policy:
        requirements = requirements.filter(policy_document__title__icontains=policy)
    if title:
        requirements = requirements.filter(title__icontains=title)
    
    # Calculate enriched context data
    mandatory_requirements = requirements.filter(mandatory=True).count()
    frameworks_count = requirements.values('regulatory_framework').distinct().count()
    jurisdictions_count = requirements.values('jurisdiction').distinct().count()
    
    # Framework distribution analysis
    framework_distribution = []
    for req in requirements.values('regulatory_framework__name').annotate(
        total_requirements=Count('id'),
        mandatory_count=Count('id', filter=Q(mandatory=True)),
        optional_count=Count('id', filter=Q(mandatory=False))
    ):
        total = req['total_requirements']
        mandatory = req['mandatory_count']
        optional = req['optional_count']
        coverage_percentage = (mandatory / total * 100) if total > 0 else 0
        
        framework_distribution.append({
            'name': req['regulatory_framework__name'] or 'Internal Requirements',
            'total_requirements': total,
            'mandatory_count': mandatory,
            'optional_count': optional,
            'coverage_percentage': coverage_percentage
        })
    
    # Jurisdiction distribution analysis
    jurisdiction_distribution = []
    for req in requirements.values('jurisdiction').annotate(
        total_requirements=Count('id'),
        mandatory_count=Count('id', filter=Q(mandatory=True)),
        optional_count=Count('id', filter=Q(mandatory=False))
    ):
        jurisdiction_distribution.append({
            'name': req['jurisdiction'],
            'total_requirements': req['total_requirements'],
            'mandatory_count': req['mandatory_count'],
            'optional_count': req['optional_count']
        })
    
    # Calculate additional metrics
    policy_coverage_percentage = (requirements.exclude(policy_document__isnull=True).count() / requirements.count() * 100) if requirements.count() > 0 else 0
    high_risk_jurisdictions = sum(1 for j in jurisdiction_distribution if j['mandatory_count'] > 10)
    framework_gaps = sum(1 for f in framework_distribution if f['mandatory_count'] == 0)
    policy_gaps = requirements.filter(policy_document__isnull=True).count()
    
    html_string = render_to_string('reports/compliance_requirement_summary.html', {
        'organization': org,
        'requirements': requirements,
        'filters': {'framework': framework, 'jurisdiction': jurisdiction, 'mandatory': mandatory, 'policy': policy, 'title': title},
        'mandatory_requirements': mandatory_requirements,
        'frameworks_count': frameworks_count,
        'jurisdictions_count': jurisdictions_count,
        'framework_distribution': framework_distribution,
        'jurisdiction_distribution': jurisdiction_distribution,
        'policy_coverage_percentage': policy_coverage_percentage,
        'high_risk_jurisdictions': high_risk_jurisdictions,
        'framework_gaps': framework_gaps,
        'policy_gaps': policy_gaps,
        'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Compliance Requirement Summary Report',
        'description': 'Comprehensive analysis of compliance requirements across regulatory frameworks and jurisdictions',
        'filters_summary': f"Framework: {framework or 'All'}, Jurisdiction: {jurisdiction or 'All'}, Mandatory: {mandatory or 'All'}"
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_compliance_requirement_summary.pdf"'
    return response

def compliance_obligation_register_pdf(request):
    org = request.tenant
    obligations = ComplianceObligation.objects.filter(organization=org)
    status = request.GET.get('status')
    owner = request.GET.get('owner')
    due_date = request.GET.get('due_date')
    priority = request.GET.get('priority')
    requirement = request.GET.get('requirement')
    obligation_id = request.GET.get('obligation_id')
    if status:
        obligations = obligations.filter(status=status)
    if owner:
        obligations = obligations.filter(owner__email__icontains=owner)
    if due_date:
        obligations = obligations.filter(due_date=due_date)
    if priority:
        obligations = obligations.filter(priority=priority)
    if requirement:
        obligations = obligations.filter(requirement__title__icontains=requirement)
    if obligation_id:
        obligations = obligations.filter(obligation_id__icontains=obligation_id)
    
    # Calculate enriched context data
    open_obligations = obligations.filter(status='open').count()
    overdue_obligations = obligations.filter(due_date__lt=timezone.now().date()).count()
    high_priority_obligations = obligations.filter(priority__gte=4).count()
    
    # Status distribution analysis
    status_distribution = []
    for obl in obligations.values('status').annotate(
        count=Count('id'),
        overdue_count=Count('id', filter=Q(due_date__lt=timezone.now().date()))
    ):
        total = obl['count']
        percentage = (total / obligations.count() * 100) if obligations.count() > 0 else 0
        
        status_distribution.append({
            'name': obl['status'].title(),
            'count': total,
            'percentage': percentage,
            'overdue_count': obl['overdue_count']
        })
    
    # Priority distribution analysis
    priority_distribution = []
    for obl in obligations.values('priority').annotate(
        count=Count('id'),
        overdue_count=Count('id', filter=Q(due_date__lt=timezone.now().date()))
    ):
        total = obl['count']
        percentage = (total / obligations.count() * 100) if obligations.count() > 0 else 0
        
        priority_distribution.append({
            'level': obl['priority'],
            'count': total,
            'percentage': percentage,
            'overdue_count': obl['overdue_count']
        })
    
    # Calculate additional metrics
    completion_rate = (obligations.filter(status='completed').count() / obligations.count() * 100) if obligations.count() > 0 else 0
    unassigned_obligations = obligations.filter(owner__isnull=True, owner_email__isnull=True).count()
    evidence_required_count = obligations.filter(evidence_required=True).count()
    
    html_string = render_to_string('reports/compliance_obligation_register.html', {
        'organization': org,
        'obligations': obligations,
        'filters': {'status': status, 'owner': owner, 'due_date': due_date, 'priority': priority, 'requirement': requirement, 'obligation_id': obligation_id},
        'open_obligations': open_obligations,
        'overdue_obligations': overdue_obligations,
        'high_priority_obligations': high_priority_obligations,
        'status_distribution': status_distribution,
        'priority_distribution': priority_distribution,
        'completion_rate': completion_rate,
        'unassigned_obligations': unassigned_obligations,
        'evidence_required_count': evidence_required_count,
        'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Compliance Obligation Register Report',
        'description': 'Comprehensive analysis of compliance obligations and their management status',
        'filters_summary': f"Status: {status or 'All'}, Owner: {owner or 'All'}, Priority: {priority or 'All'}"
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_compliance_obligation_register.pdf"'
    return response

def compliance_evidence_register_pdf(request):
    org = request.tenant
    evidences = ComplianceEvidence.objects.filter(organization=org)
    obligation = request.GET.get('obligation')
    document = request.GET.get('document')
    validity_start = request.GET.get('validity_start')
    validity_end = request.GET.get('validity_end')
    if obligation:
        evidences = evidences.filter(obligation__obligation_id__icontains=obligation)
    if document:
        evidences = evidences.filter(document__title__icontains=document)
    if validity_start:
        evidences = evidences.filter(validity_start__gte=validity_start)
    if validity_end:
        evidences = evidences.filter(validity_end__lte=validity_end)
    
    # Calculate enriched context data
    today = timezone.now().date()
    thirty_days_from_now = today + timedelta(days=30)
    
    valid_evidence = evidences.filter(validity_end__gte=today).count()
    expiring_evidence = evidences.filter(validity_end__lte=thirty_days_from_now, validity_end__gte=today).count()
    expired_evidence = evidences.filter(validity_end__lt=today).count()
    
    # Evidence status distribution analysis
    evidence_status_distribution = []
    total_evidences = evidences.count()
    
    # Valid evidence
    valid_count = valid_evidence
    evidence_status_distribution.append({
        'name': 'Valid',
        'count': valid_count,
        'percentage': (valid_count / total_evidences * 100) if total_evidences > 0 else 0,
        'obligations_count': evidences.filter(validity_end__gte=today).values('obligation').distinct().count()
    })
    
    # Expiring evidence
    expiring_count = expiring_evidence
    evidence_status_distribution.append({
        'name': 'Expiring',
        'count': expiring_count,
        'percentage': (expiring_count / total_evidences * 100) if total_evidences > 0 else 0,
        'obligations_count': evidences.filter(validity_end__lte=thirty_days_from_now, validity_end__gte=today).values('obligation').distinct().count()
    })
    
    # Expired evidence
    expired_count = expired_evidence
    evidence_status_distribution.append({
        'name': 'Expired',
        'count': expired_count,
        'percentage': (expired_count / total_evidences * 100) if total_evidences > 0 else 0,
        'obligations_count': evidences.filter(validity_end__lt=today).values('obligation').distinct().count()
    })
    
    # Obligation evidence coverage analysis
    obligation_evidence_coverage = []
    for obl in evidences.values('obligation__obligation_id', 'obligation__requirement__title').annotate(
        evidence_count=Count('id'),
        valid_evidence_count=Count('id', filter=Q(validity_end__gte=today))
    ):
        obligation_evidence_coverage.append({
            'obligation_id': obl['obligation__obligation_id'],
            'requirement_title': obl['obligation__requirement__title'],
            'evidence_count': obl['evidence_count'],
            'valid_evidence_count': obl['valid_evidence_count']
        })
    
    # Calculate additional metrics
    evidence_coverage_percentage = (valid_evidence / evidences.count() * 100) if evidences.count() > 0 else 0
    coverage_gaps = evidences.values('obligation').distinct().count() - len([o for o in obligation_evidence_coverage if o['valid_evidence_count'] > 0])
    single_evidence_obligations = len([o for o in obligation_evidence_coverage if o['evidence_count'] == 1])
    
    html_string = render_to_string('reports/compliance_evidence_register.html', {
        'organization': org,
        'evidences': evidences,
        'filters': {'obligation': obligation, 'document': document, 'validity_start': validity_start, 'validity_end': validity_end},
        'valid_evidence': valid_evidence,
        'expiring_evidence': expiring_evidence,
        'expired_evidence': expired_evidence,
        'evidence_status_distribution': evidence_status_distribution,
        'obligation_evidence_coverage': obligation_evidence_coverage,
        'evidence_coverage_percentage': evidence_coverage_percentage,
        'coverage_gaps': coverage_gaps,
        'single_evidence_obligations': single_evidence_obligations,
        'today': today,
        'thirty_days_from_now': thirty_days_from_now,
        'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Compliance Evidence Register Report',
        'description': 'Comprehensive analysis of compliance evidence and its validity status',
        'filters_summary': f"Obligation: {obligation or 'All'}, Document: {document or 'All'}"
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_compliance_evidence_register.pdf"'
    return response

def policy_document_register_pdf(request):
    org = request.tenant
    documents = PolicyDocument.objects.filter(organization=org)
    owner = request.GET.get('owner')
    effective_date = request.GET.get('effective_date')
    expiration_date = request.GET.get('expiration_date')
    title = request.GET.get('title')
    if owner:
        documents = documents.filter(owner__email__icontains=owner)
    if effective_date:
        documents = documents.filter(effective_date=effective_date)
    if expiration_date:
        documents = documents.filter(expiration_date=expiration_date)
    if title:
        documents = documents.filter(title__icontains=title)
    
    # Calculate enriched context data
    today = timezone.now().date()
    thirty_days_from_now = today + timedelta(days=30)
    
    active_policies = documents.filter(expiration_date__isnull=True) | documents.filter(expiration_date__gte=today)
    expiring_policies = documents.filter(expiration_date__lte=thirty_days_from_now, expiration_date__gte=today)
    expired_policies = documents.filter(expiration_date__lt=today)
    
    # Policy status distribution analysis
    policy_status_distribution = []
    total_policies = documents.count()
    
    # Active policies
    active_count = active_policies.count()
    policy_status_distribution.append({
        'name': 'Active',
        'count': active_count,
        'percentage': (active_count / total_policies * 100) if total_policies > 0 else 0,
        'owners_count': active_policies.values('owner').distinct().count()
    })
    
    # Expiring policies
    expiring_count = expiring_policies.count()
    policy_status_distribution.append({
        'name': 'Expiring',
        'count': expiring_count,
        'percentage': (expiring_count / total_policies * 100) if total_policies > 0 else 0,
        'owners_count': expiring_policies.values('owner').distinct().count()
    })
    
    # Expired policies
    expired_count = expired_policies.count()
    policy_status_distribution.append({
        'name': 'Expired',
        'count': expired_count,
        'percentage': (expired_count / total_policies * 100) if total_policies > 0 else 0,
        'owners_count': expired_policies.values('owner').distinct().count()
    })
    
    # Owner analysis
    owner_analysis = []
    for doc in documents.values('owner__email', 'owner__first_name', 'owner__last_name').annotate(
        total_policies=Count('id'),
        active_policies=Count('id', filter=Q(expiration_date__isnull=True) | Q(expiration_date__gte=today)),
        expired_policies=Count('id', filter=Q(expiration_date__lt=today))
    ):
        owner_name = f"{doc['owner__first_name']} {doc['owner__last_name']}" if doc['owner__first_name'] and doc['owner__last_name'] else doc['owner__email']
        owner_analysis.append({
            'name': owner_name,
            'total_policies': doc['total_policies'],
            'active_policies': doc['active_policies'],
            'expired_policies': doc['expired_policies']
        })
    
    # Calculate additional metrics
    policy_coverage_percentage = (active_count / total_policies * 100) if total_policies > 0 else 0
    unassigned_policies = documents.filter(owner__isnull=True, owner_email__isnull=True).count()
    outdated_versions = documents.filter(version__lt='2.0').count()  # Assuming version 2.0+ is current
    
    html_string = render_to_string('reports/policy_document_register.html', {
        'organization': org,
        'documents': documents,
        'filters': {'owner': owner, 'effective_date': effective_date, 'expiration_date': expiration_date, 'title': title},
        'active_policies': active_count,
        'expiring_policies': expiring_count,
        'expired_policies': expired_count,
        'policy_status_distribution': policy_status_distribution,
        'owner_analysis': owner_analysis,
        'policy_coverage_percentage': policy_coverage_percentage,
        'unassigned_policies': unassigned_policies,
        'outdated_versions': outdated_versions,
        'today': today,
        'thirty_days_from_now': thirty_days_from_now,
        'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Policy Document Register Report',
        'description': 'Comprehensive analysis of policy documents and their management status',
        'filters_summary': f"Owner: {owner or 'All'}, Title: {title or 'All'}"
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_policy_document_register.pdf"'
    return response

def compliance_requirement_details_pdf(request):
    org = request.tenant
    requirement = None
    obligations = ComplianceObligation.objects.none()  # Initialize as empty QuerySet
    evidences = ComplianceEvidence.objects.none()  # Initialize as empty QuerySet
    title = request.GET.get('title')
    if title:
        requirement = ComplianceRequirement.objects.filter(organization=org, title__icontains=title).first()
        if requirement:
            obligations = ComplianceObligation.objects.filter(requirement=requirement)
            evidences = ComplianceEvidence.objects.filter(obligation__in=obligations)
    
    # Create a dictionary to map obligations to their evidences for easy template access
    obligation_evidences_map = {}
    if obligations.exists():
        for obligation in obligations:
            obligation_evidences_map[obligation.id] = evidences.filter(obligation=obligation)
    
    # Calculate enriched context data
    open_obligations = obligations.filter(status='open').count()
    completed_obligations = obligations.filter(status='completed').count()
    overdue_obligations = obligations.filter(due_date__lt=timezone.now().date()).count()
    high_priority_obligations = obligations.filter(priority__gte=4).count()
    
    # Obligation status distribution analysis
    obligation_status_distribution = []
    if obligations.exists():
        for obl in obligations.values('status').annotate(
            count=Count('id'),
            overdue_count=Count('id', filter=Q(due_date__lt=timezone.now().date()))
        ):
            total = obl['count']
            percentage = (total / obligations.count() * 100) if obligations.count() > 0 else 0
            
            obligation_status_distribution.append({
                'name': obl['status'].title(),
                'count': total,
                'percentage': percentage,
                'overdue_count': obl['overdue_count']
            })
    
    # Calculate additional metrics
    compliance_rate = (completed_obligations / obligations.count() * 100) if obligations.count() > 0 else 0
    evidence_gaps = obligations.filter(evidence_required=True).exclude(complianceevidence__isnull=False).count()
    unassigned_obligations = obligations.filter(owner__isnull=True, owner_email__isnull=True).count()
    
    html_string = render_to_string('reports/compliance_requirement_details.html', {
        'organization': org,
        'requirement': requirement,
        'obligations': obligations,
        'evidences': evidences,
        'obligation_evidences_map': obligation_evidences_map,
        'open_obligations': open_obligations,
        'completed_obligations': completed_obligations,
        'overdue_obligations': overdue_obligations,
        'high_priority_obligations': high_priority_obligations,
        'obligation_status_distribution': obligation_status_distribution,
        'compliance_rate': compliance_rate,
        'evidence_gaps': evidence_gaps,
        'unassigned_obligations': unassigned_obligations,
        'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Compliance Requirement Details Report',
        'description': 'Detailed analysis of compliance requirement implementation and associated obligations',
        'filters_summary': f"Requirement: {title or 'Not specified'}"
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_compliance_requirement_details.pdf"'
    return response

def compliance_obligation_details_pdf(request):
    org = request.tenant
    obligation = None
    evidences = ComplianceEvidence.objects.none()  # Initialize as empty QuerySet
    obligation_id = request.GET.get('obligation_id')
    if obligation_id:
        obligation = ComplianceObligation.objects.filter(organization=org, obligation_id__icontains=obligation_id).first()
        if obligation:
            evidences = ComplianceEvidence.objects.filter(obligation=obligation)
    
    # Calculate enriched context data
    today = timezone.now().date()
    thirty_days_from_now = today + timedelta(days=30)
    
    valid_evidence = evidences.filter(validity_end__gte=today).count()
    expiring_evidence = evidences.filter(validity_end__lte=thirty_days_from_now, validity_end__gte=today).count()
    expired_evidence = evidences.filter(validity_end__lt=today).count()
    
    # Evidence status distribution analysis
    evidence_status_distribution = []
    if evidences.exists():
        total_evidences = evidences.count()
        
        # Valid evidence
        valid_count = valid_evidence
        evidence_status_distribution.append({
            'name': 'Valid',
            'count': valid_count,
            'percentage': (valid_count / total_evidences * 100) if total_evidences > 0 else 0,
            'documents_count': evidences.filter(validity_end__gte=today).values('document').distinct().count()
        })
        
        # Expiring evidence
        expiring_count = expiring_evidence
        evidence_status_distribution.append({
            'name': 'Expiring',
            'count': expiring_count,
            'percentage': (expiring_count / total_evidences * 100) if total_evidences > 0 else 0,
            'documents_count': evidences.filter(validity_end__lte=thirty_days_from_now, validity_end__gte=today).values('document').distinct().count()
        })
        
        # Expired evidence
        expired_count = expired_evidence
        evidence_status_distribution.append({
            'name': 'Expired',
            'count': expired_count,
            'percentage': (expired_count / total_evidences * 100) if total_evidences > 0 else 0,
            'documents_count': evidences.filter(validity_end__lt=today).values('document').distinct().count()
        })
    
    # Calculate risk levels and descriptions
    if obligation:
        # Obligation risk assessment
        if obligation.check_overdue:
            obligation_risk_level = "High Risk"
            obligation_risk_description = "Obligation is overdue and requires immediate attention"
        elif obligation.status == 'open' and obligation.due_date <= today:
            obligation_risk_level = "Medium Risk"
            obligation_risk_description = "Obligation is due today and should be completed"
        else:
            obligation_risk_level = "Low Risk"
            obligation_risk_description = "Obligation is on track"
        
        # Priority description
        if obligation.priority >= 4:
            priority_description = "Critical priority requiring immediate attention"
        elif obligation.priority >= 3:
            priority_description = "High priority requiring prompt action"
        else:
            priority_description = "Standard priority"
        
        # Evidence risk assessment
        if expired_evidence > 0:
            evidence_risk_level = "High Risk"
            evidence_risk_description = "Evidence has expired and needs renewal"
        elif expiring_evidence > 0:
            evidence_risk_level = "Medium Risk"
            evidence_risk_description = "Evidence is expiring soon and needs attention"
        else:
            evidence_risk_level = "Low Risk"
            evidence_risk_description = "Evidence is valid and current"
        
        # Compliance risk assessment
        if obligation.check_overdue or expired_evidence > 0:
            compliance_risk_level = "High Risk"
            compliance_risk_description = "Compliance risk due to overdue obligation or expired evidence"
        elif obligation.status == 'open' and obligation.due_date <= today or expiring_evidence > 0:
            compliance_risk_level = "Medium Risk"
            compliance_risk_description = "Compliance risk due to due obligation or expiring evidence"
        else:
            compliance_risk_level = "Low Risk"
            compliance_risk_description = "Compliance is on track"
        
        # Owner accountability status
        if obligation.owner or obligation.owner_email:
            owner_accountability_status = "Owner assigned and accountable"
        else:
            owner_accountability_status = "No owner assigned - requires attention"
        
        # Recommendations
        obligation_management_recommendation = "Ensure obligation is properly managed and tracked"
        evidence_collection_recommendation = f"Collect evidence for {obligation.evidence_required and evidences.count() == 0 and 'this obligation' or 'all obligations'}"
        status_update_recommendation = "Regularly update obligation status to maintain accurate tracking"
        owner_engagement_recommendation = "Engage with obligation owner to ensure timely completion"
        risk_mitigation_recommendation = "Implement risk mitigation strategies for high-priority obligations"
        compliance_monitoring_recommendation = "Establish regular compliance monitoring and reporting"
    else:
        obligation_risk_level = "N/A"
        obligation_risk_description = "No obligation found"
        priority_description = "N/A"
        evidence_risk_level = "N/A"
        evidence_risk_description = "No evidence found"
        compliance_risk_level = "N/A"
        compliance_risk_description = "No compliance data available"
        owner_accountability_status = "N/A"
        obligation_management_recommendation = "N/A"
        evidence_collection_recommendation = "N/A"
        status_update_recommendation = "N/A"
        owner_engagement_recommendation = "N/A"
        risk_mitigation_recommendation = "N/A"
        compliance_monitoring_recommendation = "N/A"
    
    html_string = render_to_string('reports/compliance_obligation_details.html', {
        'organization': org,
        'obligation': obligation,
        'evidences': evidences,
        'valid_evidence': valid_evidence,
        'expiring_evidence': expiring_evidence,
        'expired_evidence': expired_evidence,
        'evidence_status_distribution': evidence_status_distribution,
        'obligation_risk_level': obligation_risk_level,
        'obligation_risk_description': obligation_risk_description,
        'priority_description': priority_description,
        'evidence_risk_level': evidence_risk_level,
        'evidence_risk_description': evidence_risk_description,
        'compliance_risk_level': compliance_risk_level,
        'compliance_risk_description': compliance_risk_description,
        'owner_accountability_status': owner_accountability_status,
        'obligation_management_recommendation': obligation_management_recommendation,
        'evidence_collection_recommendation': evidence_collection_recommendation,
        'status_update_recommendation': status_update_recommendation,
        'owner_engagement_recommendation': owner_engagement_recommendation,
        'risk_mitigation_recommendation': risk_mitigation_recommendation,
        'compliance_monitoring_recommendation': compliance_monitoring_recommendation,
        'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Compliance Obligation Details Report',
        'description': 'Detailed analysis of compliance obligation and associated evidence',
        'filters_summary': f"Obligation ID: {obligation_id or 'Not specified'}"
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm 1cm }')])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_compliance_obligation_details.pdf"'
    return response

def contract_register_summary_pdf(request):
    org = request.tenant
    contracts = Contract.objects.filter(organization=org)
    status = request.GET.get('status')
    type_ = request.GET.get('type')
    party = request.GET.get('party')
    title = request.GET.get('title')
    if status:
        contracts = contracts.filter(status=status)
    if type_:
        contracts = contracts.filter(contract_type__name__icontains=type_)
    if party:
        contracts = contracts.filter(parties__name__icontains=party)
    if title:
        contracts = contracts.filter(title__icontains=title)
    filters = {'status': status, 'type': type_, 'party': party, 'title': title}
    return render(request, 'reports/contract_register_summary.html', {
        'contracts': contracts,
        'organization': org,
        'filters': filters,
    })

def contract_register_detailed_pdf(request):
    org = request.tenant
    contracts = Contract.objects.filter(organization=org)
    status = request.GET.get('status')
    type_ = request.GET.get('type')
    party = request.GET.get('party')
    title = request.GET.get('title')
    if status:
        contracts = contracts.filter(status=status)
    if type_:
        contracts = contracts.filter(contract_type__name__icontains=type_)
    if party:
        contracts = contracts.filter(parties__name__icontains=party)
    if title:
        contracts = contracts.filter(title__icontains=title)
    filters = {'status': status, 'type': type_, 'party': party, 'title': title}
    return render(request, 'reports/contract_register_detailed.html', {
        'contracts': contracts,
        'organization': org,
        'filters': filters,
    })

def milestone_register_pdf(request):
    org = request.tenant
    milestones = ContractMilestone.objects.filter(organization=org)
    type_ = request.GET.get('type')
    due_date = request.GET.get('due_date')
    status = request.GET.get('status')
    if type_:
        milestones = milestones.filter(milestone_type__icontains=type_)
    if due_date:
        milestones = milestones.filter(due_date=due_date)
    if status:
        milestones = milestones.filter(status=status)
    filters = {'type': type_, 'due_date': due_date, 'status': status}
    return render(request, 'reports/milestone_register.html', {
        'milestones': milestones,
        'organization': org,
        'filters': filters,
    })

def party_register_pdf(request):
    org = request.tenant
    parties = Party.objects.filter(organization=org)
    party_name = request.GET.get('party_name')
    role = request.GET.get('role')
    if party_name:
        parties = parties.filter(name__icontains=party_name)
    if role:
        parties = parties.filter(role__icontains=role)
    filters = {'party_name': party_name, 'role': role}
    return render(request, 'reports/party_register.html', {
        'parties': parties,
        'organization': org,
        'filters': filters,
    })

def contract_expiry_pdf(request):
    org = request.tenant
    contracts = Contract.objects.filter(organization=org)
    # Optionally filter by expiry date range
    expiry_start = request.GET.get('expiry_start')
    expiry_end = request.GET.get('expiry_end')
    if expiry_start:
        contracts = contracts.filter(expiry_date__gte=expiry_start)
    if expiry_end:
        contracts = contracts.filter(expiry_date__lte=expiry_end)
    filters = {'expiry_start': expiry_start, 'expiry_end': expiry_end}
    return render(request, 'reports/contract_expiry.html', {
        'contracts': contracts,
        'organization': org,
        'filters': filters,
    })

def contract_details_pdf(request):
    org = request.tenant
    contracts = Contract.objects.filter(organization=org)
    return render(request, 'reports/contract_details.html', {
        'contracts': contracts,
        'organization': org
    })

def milestone_details_pdf(request):
    org = request.tenant
    milestones = ContractMilestone.objects.filter(organization=org)
    return render(request, 'reports/milestone_details.html', {
        'milestones': milestones,
        'organization': org
    })

def party_details_pdf(request):
    """Generate party details report with optional filtering by party_id"""
    org = request.tenant
    party_id = request.GET.get('party_id')
    parties = Party.objects.filter(organization=org)
    if party_id:
        parties = parties.filter(id=party_id)

    html_string = render_to_string('reports/party_details.html', {
        'organization': org,
        'parties': parties,
        'filters': {'party_id': party_id},
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'Party Details Report',
        'description': 'Detailed contract parties analysis with related contracts',
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_party_details.pdf"'
    return response

# ─── CONTRACT REPORTS ────────────────────────────────────────────────────────────

def contract_register_summary_pdf(request):
    """Generate comprehensive contract register summary report"""
    org = request.tenant
    contracts = Contract.objects.filter(organization=org)
    
    # Apply filters
    status = request.GET.get('status')
    contract_type = request.GET.get('contract_type')
    party = request.GET.get('party')
    title = request.GET.get('title')
    
    if status:
        contracts = contracts.filter(status=status)
    if contract_type:
        contracts = contracts.filter(contract_type__name__icontains=contract_type)
    if party:
        contracts = contracts.filter(parties__name__icontains=party)
    if title:
        contracts = contracts.filter(title__icontains=title)
    
    # Calculate summary statistics
    total_contracts = contracts.count()
    active_contracts = contracts.filter(status='active').count()
    expired_contracts = contracts.filter(status='expired').count()
    draft_contracts = contracts.filter(status='draft').count()
    terminated_contracts = contracts.filter(status='terminated').count()
    
    # Status distribution
    status_distribution = contracts.values('status').annotate(count=Count('id'))
    
    # Contract type distribution
    type_distribution = contracts.values('contract_type__name').annotate(count=Count('id'))
    
    # Party distribution
    party_distribution = contracts.values('parties__name').annotate(count=Count('id')).exclude(parties__name__isnull=True)
    
    # Expiry analysis
    today = timezone.now().date()
    expiring_30_days = contracts.filter(end_date__lte=today + timedelta(days=30), end_date__gte=today).count()
    expiring_90_days = contracts.filter(end_date__lte=today + timedelta(days=90), end_date__gte=today).count()
    
    # Financial analysis
    total_value = contracts.aggregate(total=models.Sum('value'))['total'] or 0
    
    filters = {'status': status, 'contract_type': contract_type, 'party': party, 'title': title}
    
    html_string = render_to_string('reports/contract_register_summary.html', {
        'organization': org,
        'contracts': contracts,
        'total_contracts': total_contracts,
        'active_contracts': active_contracts,
        'expired_contracts': expired_contracts,
        'draft_contracts': draft_contracts,
        'terminated_contracts': terminated_contracts,
        'status_distribution': status_distribution,
        'type_distribution': type_distribution,
        'party_distribution': party_distribution,
        'expiring_30_days': expiring_30_days,
        'expiring_90_days': expiring_90_days,
        'total_value': total_value,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_contract_register_summary.pdf"'
    return response

def contract_register_detailed_pdf(request):
    """Generate detailed contract register report"""
    org = request.tenant
    contracts = Contract.objects.filter(organization=org)
    
    # Apply filters
    status = request.GET.get('status')
    contract_type = request.GET.get('contract_type')
    party = request.GET.get('party')
    title = request.GET.get('title')
    
    if status:
        contracts = contracts.filter(status=status)
    if contract_type:
        contracts = contracts.filter(contract_type__name__icontains=contract_type)
    if party:
        contracts = contracts.filter(parties__name__icontains=party)
    if title:
        contracts = contracts.filter(title__icontains=title)
    
    # Calculate detailed statistics
    total_contracts = contracts.count()
    active_contracts = contracts.filter(status='active').count()
    expired_contracts = contracts.filter(status='expired').count()
    draft_contracts = contracts.filter(status='draft').count()
    terminated_contracts = contracts.filter(status='terminated').count()
    
    # Financial analysis
    total_value = contracts.aggregate(total=models.Sum('value'))['total'] or 0
    avg_value = contracts.aggregate(avg=models.Avg('value'))['avg'] or 0
    
    # Timeline analysis
    today = timezone.now().date()
    expiring_30_days = contracts.filter(end_date__lte=today + timedelta(days=30), end_date__gte=today).count()
    expiring_90_days = contracts.filter(end_date__lte=today + timedelta(days=90), end_date__gte=today).count()
    overdue_contracts = contracts.filter(end_date__lt=today, status='active').count()
    
    # Compliance analysis
    contracts_with_obligations = contracts.filter(compliance_obligations__isnull=False).distinct().count()
    
    filters = {'status': status, 'contract_type': contract_type, 'party': party, 'title': title}
    
    html_string = render_to_string('reports/contract_register_detailed.html', {
        'organization': org,
        'contracts': contracts,
        'total_contracts': total_contracts,
        'active_contracts': active_contracts,
        'expired_contracts': expired_contracts,
        'draft_contracts': draft_contracts,
        'terminated_contracts': terminated_contracts,
        'total_value': total_value,
        'avg_value': avg_value,
        'expiring_30_days': expiring_30_days,
        'expiring_90_days': expiring_90_days,
        'overdue_contracts': overdue_contracts,
        'contracts_with_obligations': contracts_with_obligations,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_contract_register_detailed.pdf"'
    return response

def milestone_register_pdf(request):
    """Generate milestone register report"""
    org = request.tenant
    milestones = ContractMilestone.objects.filter(organization=org)
    
    # Apply filters
    milestone_type = request.GET.get('milestone_type')
    is_completed = request.GET.get('is_completed')
    contract = request.GET.get('contract')
    
    if milestone_type:
        milestones = milestones.filter(milestone_type=milestone_type)
    if is_completed:
        milestones = milestones.filter(is_completed=is_completed.lower() == 'true')
    if contract:
        milestones = milestones.filter(contract__title__icontains=contract)
    
    # Calculate statistics
    total_milestones = milestones.count()
    completed_milestones = milestones.filter(is_completed=True).count()
    pending_milestones = milestones.filter(is_completed=False).count()
    
    # Overdue analysis
    today = timezone.now().date()
    overdue_milestones = milestones.filter(due_date__lt=today, is_completed=False).count()
    due_soon_milestones = milestones.filter(
        due_date__gte=today, 
        due_date__lte=today + timedelta(days=7), 
        is_completed=False
    ).count()
    
    # Type distribution
    type_distribution = milestones.values('milestone_type').annotate(count=Count('id'))
    
    # Contract distribution
    contract_distribution = milestones.values('contract__title').annotate(count=Count('id'))
    
    filters = {'milestone_type': milestone_type, 'is_completed': is_completed, 'contract': contract}
    
    html_string = render_to_string('reports/milestone_register.html', {
        'organization': org,
        'milestones': milestones,
        'total_milestones': total_milestones,
        'completed_milestones': completed_milestones,
        'pending_milestones': pending_milestones,
        'overdue_milestones': overdue_milestones,
        'due_soon_milestones': due_soon_milestones,
        'type_distribution': type_distribution,
        'contract_distribution': contract_distribution,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_milestone_register.pdf"'
    return response

def party_register_pdf(request):
    """Generate party register report"""
    org = request.tenant
    parties = Party.objects.filter(organization=org)
    
    # Apply filters
    party_type = request.GET.get('party_type')
    name = request.GET.get('name')
    
    if party_type:
        parties = parties.filter(party_type=party_type)
    if name:
        parties = parties.filter(name__icontains=name)
    
    # Calculate statistics
    total_parties = parties.count()
    internal_parties = parties.filter(party_type='internal').count()
    external_parties = parties.filter(party_type='external').count()
    government_parties = parties.filter(party_type='government').count()
    third_party_parties = parties.filter(party_type='third_party').count()
    
    # Type distribution
    type_distribution = parties.values('party_type').annotate(count=Count('id'))
    
    # Contract participation analysis
    parties_with_contracts = parties.annotate(
        contract_count=Count('contracts', filter=Q(contracts__organization=org))
    ).filter(contract_count__gt=0)
    
    filters = {'party_type': party_type, 'name': name}
    
    html_string = render_to_string('reports/party_register.html', {
        'organization': org,
        'parties': parties,
        'total_parties': total_parties,
        'internal_parties': internal_parties,
        'external_parties': external_parties,
        'government_parties': government_parties,
        'third_party_parties': third_party_parties,
        'type_distribution': type_distribution,
        'parties_with_contracts': parties_with_contracts,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_party_register.pdf"'
    return response

def contract_expiry_pdf(request):
    """Generate contract expiry analysis report"""
    org = request.tenant
    contracts = Contract.objects.filter(organization=org)
    
    # Apply filters
    expiry_start = request.GET.get('expiry_start')
    expiry_end = request.GET.get('expiry_end')
    
    if expiry_start:
        contracts = contracts.filter(end_date__gte=expiry_start)
    if expiry_end:
        contracts = contracts.filter(end_date__lte=expiry_end)
    
    # Expiry analysis
    today = timezone.now().date()
    expired_contracts = contracts.filter(end_date__lt=today)
    expiring_30_days = contracts.filter(end_date__lte=today + timedelta(days=30), end_date__gte=today)
    expiring_90_days = contracts.filter(end_date__lte=today + timedelta(days=90), end_date__gte=today)
    valid_contracts = contracts.filter(end_date__gt=today + timedelta(days=90))
    
    # Financial impact analysis
    expired_value = expired_contracts.aggregate(total=models.Sum('value'))['total'] or 0
    expiring_30_value = expiring_30_days.aggregate(total=models.Sum('value'))['total'] or 0
    expiring_90_value = expiring_90_days.aggregate(total=models.Sum('value'))['total'] or 0
    
    # Auto-renewal analysis
    auto_renew_contracts = contracts.filter(auto_renew=True)
    auto_renew_expiring = auto_renew_contracts.filter(end_date__lte=today + timedelta(days=90))
    
    filters = {'expiry_start': expiry_start, 'expiry_end': expiry_end}
    
    html_string = render_to_string('reports/contract_expiry.html', {
        'organization': org,
        'expired_contracts': expired_contracts,
        'expiring_30_days': expiring_30_days,
        'expiring_90_days': expiring_90_days,
        'valid_contracts': valid_contracts,
        'expired_value': expired_value,
        'expiring_30_value': expiring_30_value,
        'expiring_90_value': expiring_90_value,
        'auto_renew_contracts': auto_renew_contracts,
        'auto_renew_expiring': auto_renew_expiring,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_contract_expiry.pdf"'
    return response

def contract_details_pdf(request):
    """Generate detailed contract analysis report"""
    org = request.tenant
    contract_id = request.GET.get('contract_id')
    contract = None
    
    if contract_id:
        contract = Contract.objects.filter(organization=org, id=contract_id).first()
    
    if not contract:
        # If no specific contract, show the most recent one
        contract = Contract.objects.filter(organization=org).order_by('-created_at').first()
    
    if contract:
        # Get related data
        milestones = contract.milestones.all()
        parties = contract.parties.all()
        obligations = contract.compliance_obligations.all()
        
        # Milestone analysis
        total_milestones = milestones.count()
        completed_milestones = milestones.filter(is_completed=True).count()
        overdue_milestones = milestones.filter(due_date__lt=timezone.now().date(), is_completed=False).count()
        
        # Financial analysis
        days_to_expiry = contract.days_to_expiry
        is_overdue = contract.is_overdue()
        
        html_string = render_to_string('reports/contract_details.html', {
            'organization': org,
            'contract': contract,
            'milestones': milestones,
            'parties': parties,
            'obligations': obligations,
            'total_milestones': total_milestones,
            'completed_milestones': completed_milestones,
            'overdue_milestones': overdue_milestones,
            'days_to_expiry': days_to_expiry,
            'is_overdue': is_overdue,
            'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        })
    else:
        html_string = render_to_string('reports/contract_details.html', {
            'organization': org,
            'contract': None,
            'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_contract_details.pdf"'
    return response

def milestone_details_pdf(request):
    """Generate detailed milestone analysis report"""
    org = request.tenant
    milestone_id = request.GET.get('milestone_id')
    milestone = None
    
    if milestone_id:
        milestone = ContractMilestone.objects.filter(organization=org, id=milestone_id).first()
    
    if not milestone:
        # If no specific milestone, show the most recent one
        milestone = ContractMilestone.objects.filter(organization=org).order_by('-created_at').first()
    
    if milestone:
        # Calculate metrics
        days_until_due = (milestone.due_date - timezone.now().date()).days
        is_overdue = milestone.check_overdue()
        
        html_string = render_to_string('reports/milestone_details.html', {
            'organization': org,
            'milestone': milestone,
            'days_until_due': days_until_due,
            'is_overdue': is_overdue,
            'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        })
    else:
        html_string = render_to_string('reports/milestone_details.html', {
            'organization': org,
            'milestone': None,
            'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_milestone_details.pdf"'
    return response

# ─── LEGAL REPORTS ────────────────────────────────────────────────────────────

def legal_case_summary_pdf(request):
    """Generate comprehensive legal case summary report"""
    org = request.tenant
    cases = LegalCase.objects.filter(organization=org)
    
    # Apply filters
    case_type = request.GET.get('case_type')
    status = request.GET.get('status')
    priority = request.GET.get('priority')
    
    if case_type:
        cases = cases.filter(case_type__name__icontains=case_type)
    if status:
        cases = cases.filter(status=status)
    if priority:
        cases = cases.filter(priority=priority)
    
    # Calculate summary statistics
    total_cases = cases.count()
    open_cases = cases.filter(status__in=['intake', 'investigation', 'litigation', 'settlement_negotiation']).count()
    closed_cases = cases.filter(status='closed').count()
    archived_cases = cases.filter(status='archived').count()
    overdue_cases = sum(1 for case in cases if case.is_overdue())
    
    # Case type distribution
    case_type_distribution = cases.values('case_type__name').annotate(count=Count('id'))
    
    # Status distribution
    status_distribution = cases.values('status').annotate(count=Count('id'))
    
    # Priority distribution
    priority_distribution = cases.values('priority').annotate(count=Count('id'))
    
    # Recent cases
    recent_cases = cases.order_by('-opened_date')[:10]
    
    filters = {'case_type': case_type, 'status': status, 'priority': priority}
    
    html_string = render_to_string('reports/legal_case_summary.html', {
        'organization': org,
        'cases': cases,
        'total_cases': total_cases,
        'open_cases': open_cases,
        'closed_cases': closed_cases,
        'archived_cases': archived_cases,
        'overdue_cases': overdue_cases,
        'case_type_distribution': case_type_distribution,
        'status_distribution': status_distribution,
        'priority_distribution': priority_distribution,
        'recent_cases': recent_cases,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_case_summary.pdf"'
    return response

def legal_task_register_pdf(request):
    """Generate legal task register report"""
    org = request.tenant
    tasks = LegalTask.objects.filter(organization=org)
    
    # Apply filters
    status = request.GET.get('status')
    assigned_to = request.GET.get('assigned_to')
    due_date = request.GET.get('due_date')
    
    if status:
        tasks = tasks.filter(status=status)
    if assigned_to:
        tasks = tasks.filter(assigned_to__email__icontains=assigned_to)
    if due_date:
        tasks = tasks.filter(due_date=due_date)
    
    # Calculate summary statistics
    total_tasks = tasks.count()
    pending_tasks = tasks.filter(status='pending').count()
    in_progress_tasks = tasks.filter(status='in_progress').count()
    completed_tasks = tasks.filter(status='completed').count()
    overdue_tasks = tasks.filter(status='overdue').count()
    
    # Status distribution
    status_distribution = tasks.values('status').annotate(count=Count('id'))
    
    # Assigned attorney distribution
    attorney_distribution = tasks.values('assigned_to__email').annotate(count=Count('id'))
    
    # Due date analysis
    overdue_tasks_list = tasks.filter(status='overdue').order_by('due_date')
    upcoming_tasks = tasks.filter(due_date__gte=timezone.now().date()).order_by('due_date')[:10]
    
    filters = {'status': status, 'assigned_to': assigned_to, 'due_date': due_date}
    
    html_string = render_to_string('reports/legal_task_register.html', {
        'organization': org,
        'tasks': tasks,
        'total_tasks': total_tasks,
        'pending_tasks': pending_tasks,
        'in_progress_tasks': in_progress_tasks,
        'completed_tasks': completed_tasks,
        'overdue_tasks': overdue_tasks,
        'status_distribution': status_distribution,
        'attorney_distribution': attorney_distribution,
        'overdue_tasks_list': overdue_tasks_list,
        'upcoming_tasks': upcoming_tasks,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_task_register.pdf"'
    return response

def legal_document_register_pdf(request):
    """Generate legal document register report"""
    org = request.tenant
    documents = LegalDocument.objects.filter(organization=org)
    
    # Apply filters
    document_title = request.GET.get('document_title')
    is_confidential = request.GET.get('is_confidential')
    
    if document_title:
        documents = documents.filter(title__icontains=document_title)
    if is_confidential:
        documents = documents.filter(is_confidential=is_confidential.lower() == 'true')
    
    # Calculate summary statistics
    total_documents = documents.count()
    confidential_documents = documents.filter(is_confidential=True).count()
    non_confidential_documents = documents.filter(is_confidential=False).count()
    
    # Document distribution by case
    case_distribution = documents.values('case__title').annotate(count=Count('id'))
    
    # Version analysis
    version_distribution = documents.values('version').annotate(count=Count('id'))
    
    # Recent documents
    recent_documents = documents.order_by('-created_at')[:10]
    
    filters = {'document_title': document_title, 'is_confidential': is_confidential}
    
    html_string = render_to_string('reports/legal_document_register.html', {
        'organization': org,
        'documents': documents,
        'total_documents': total_documents,
        'confidential_documents': confidential_documents,
        'non_confidential_documents': non_confidential_documents,
        'case_distribution': case_distribution,
        'version_distribution': version_distribution,
        'recent_documents': recent_documents,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_document_register.pdf"'
    return response

def legal_party_register_pdf(request):
    """Generate legal party register report"""
    org = request.tenant
    parties = LegalParty.objects.filter(organization=org)
    
    # Apply filters
    party_type = request.GET.get('party_type')
    party_name = request.GET.get('party_name')
    
    if party_type:
        parties = parties.filter(party_type=party_type)
    if party_name:
        parties = parties.filter(name__icontains=party_name)
    
    # Calculate summary statistics
    total_parties = parties.count()
    plaintiffs = parties.filter(party_type='plaintiff').count()
    defendants = parties.filter(party_type='defendant').count()
    witnesses = parties.filter(party_type='witness').count()
    third_parties = parties.filter(party_type='third_party').count()
    
    # Party type distribution
    party_type_distribution = parties.values('party_type').annotate(count=Count('id'))
    
    # Recent parties
    recent_parties = parties.order_by('-created_at')[:10]
    
    filters = {'party_type': party_type, 'party_name': party_name}
    
    html_string = render_to_string('reports/legal_party_register.html', {
        'organization': org,
        'parties': parties,
        'total_parties': total_parties,
        'plaintiffs': plaintiffs,
        'defendants': defendants,
        'witnesses': witnesses,
        'third_parties': third_parties,
        'party_type_distribution': party_type_distribution,
        'recent_parties': recent_parties,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_party_register.pdf"'
    return response

def legal_archive_register_pdf(request):
    """Generate legal archive register report"""
    org = request.tenant
    archives = LegalArchive.objects.filter(organization=org)
    
    # Apply filters
    archive_date = request.GET.get('archive_date')
    retention_period = request.GET.get('retention_period')
    
    if archive_date:
        archives = archives.filter(archive_date=archive_date)
    if retention_period:
        archives = archives.filter(retention_period_years=retention_period)
    
    # Calculate summary statistics
    total_archives = archives.count()
    total_retention_years = sum(archive.retention_period_years for archive in archives)
    avg_retention_years = total_retention_years / total_archives if total_archives > 0 else 0
    
    # Retention period distribution
    retention_distribution = archives.values('retention_period_years').annotate(count=Count('id'))
    
    # Archive date analysis
    recent_archives = archives.order_by('-archive_date')[:10]
    
    # Destruction date analysis
    upcoming_destruction = archives.filter(destruction_date__gte=timezone.now().date()).order_by('destruction_date')[:10]
    
    filters = {'archive_date': archive_date, 'retention_period': retention_period}
    
    html_string = render_to_string('reports/legal_archive_register.html', {
        'organization': org,
        'archives': archives,
        'total_archives': total_archives,
        'total_retention_years': total_retention_years,
        'avg_retention_years': avg_retention_years,
        'retention_distribution': retention_distribution,
        'recent_archives': recent_archives,
        'upcoming_destruction': upcoming_destruction,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_archive_register.pdf"'
    return response

def legal_case_details_pdf(request):
    """Generate detailed legal case report"""
    org = request.tenant
    cases = LegalCase.objects.filter(organization=org)
    
    # Apply filters
    title = request.GET.get('title')
    
    if title:
        cases = cases.filter(title__icontains=title)
    
    # Get detailed case information
    detailed_cases = []
    for case in cases:
        case_data = {
            'case': case,
            'tasks': case.tasks.all(),
            'documents': case.documents.all(),
            'parties': case.parties.all(),
            'attorneys': case.attorneys.all(),
        }
        detailed_cases.append(case_data)
    
    filters = {'title': title}
    
    html_string = render_to_string('reports/legal_case_details.html', {
        'organization': org,
        'detailed_cases': detailed_cases,
        'total_cases': cases.count(),
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_case_details.pdf"'
    return response

def legal_task_details_pdf(request):
    """Generate detailed legal task report"""
    org = request.tenant
    tasks = LegalTask.objects.filter(organization=org)
    
    # Apply filters
    title = request.GET.get('title')
    
    if title:
        tasks = tasks.filter(title__icontains=title)
    
    # Calculate task statistics
    total_tasks = tasks.count()
    completed_tasks = tasks.filter(status='completed').count()
    pending_tasks = tasks.filter(status='pending').count()
    overdue_tasks = tasks.filter(status='overdue').count()
    
    # Task completion analysis
    completion_rate = (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0
    
    # Average completion time
    completed_tasks_with_dates = tasks.filter(status='completed', completion_date__isnull=False)
    avg_completion_days = 0
    if completed_tasks_with_dates.exists():
        total_days = sum((task.completion_date - task.due_date).days for task in completed_tasks_with_dates if task.completion_date and task.due_date)
        avg_completion_days = total_days / completed_tasks_with_dates.count()
    
    filters = {'title': title}
    
    html_string = render_to_string('reports/legal_task_details.html', {
        'organization': org,
        'tasks': tasks,
        'total_tasks': total_tasks,
        'completed_tasks': completed_tasks,
        'pending_tasks': pending_tasks,
        'overdue_tasks': overdue_tasks,
        'completion_rate': completion_rate,
        'avg_completion_days': avg_completion_days,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_legal_task_details.pdf"'
    return response

# COBIT Reports
def cobit_domain_summary_pdf(request):
    """Generate COBIT domain summary report"""
    from risk.models import COBITDomain, COBITProcess, COBITCapability, COBITControl
    
    org = request.tenant
    domains = COBITDomain.objects.filter(organization=org)
    
    # Apply filters
    domain_code = request.GET.get('domain_code')
    if domain_code:
        domains = domains.filter(domain_code=domain_code)
    
    # Calculate statistics
    total_processes = COBITProcess.objects.filter(domain__organization=org).count()
    total_capabilities = COBITCapability.objects.filter(process__domain__organization=org).count()
    total_controls = COBITControl.objects.filter(process__domain__organization=org).count()
    
    filters = {'domain_code': domain_code}
    
    html_string = render_to_string('reports/cobit_domain_summary.html', {
        'organization': org,
        'domains': domains,
        'total_processes': total_processes,
        'total_capabilities': total_capabilities,
        'total_controls': total_controls,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_cobit_domain_summary.pdf"'
    return response

def cobit_capability_assessment_pdf(request):
    """Generate COBIT capability assessment report"""
    from risk.models import COBITCapability
    
    org = request.tenant
    capabilities = COBITCapability.objects.filter(process__domain__organization=org)
    
    # Apply filters
    process = request.GET.get('process')
    current_maturity = request.GET.get('current_maturity')
    target_maturity = request.GET.get('target_maturity')
    
    if process:
        capabilities = capabilities.filter(process__process_code__icontains=process)
    if current_maturity:
        capabilities = capabilities.filter(current_maturity=current_maturity)
    if target_maturity:
        capabilities = capabilities.filter(target_maturity=target_maturity)
    
    # Calculate statistics
    avg_current_maturity = capabilities.aggregate(avg=models.Avg('current_maturity'))['avg'] or 0
    avg_target_maturity = capabilities.aggregate(avg=models.Avg('target_maturity'))['avg'] or 0
    maturity_gap = avg_target_maturity - avg_current_maturity
    
    filters = {
        'process': process,
        'current_maturity': current_maturity,
        'target_maturity': target_maturity
    }
    
    html_string = render_to_string('reports/cobit_capability_assessment.html', {
        'organization': org,
        'capabilities': capabilities,
        'avg_current_maturity': avg_current_maturity,
        'avg_target_maturity': avg_target_maturity,
        'maturity_gap': maturity_gap,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_cobit_capability_assessment.pdf"'
    return response

def cobit_control_effectiveness_pdf(request):
    """Generate COBIT control effectiveness report"""
    from risk.models import COBITControl
    
    org = request.tenant
    controls = COBITControl.objects.filter(process__domain__organization=org)
    
    # Apply filters
    process = request.GET.get('process')
    control_type = request.GET.get('control_type')
    implementation_status = request.GET.get('implementation_status')
    effectiveness_rating = request.GET.get('effectiveness_rating')
    
    if process:
        controls = controls.filter(process__process_code__icontains=process)
    if control_type:
        controls = controls.filter(control_type=control_type)
    if implementation_status:
        controls = controls.filter(implementation_status=implementation_status)
    if effectiveness_rating:
        controls = controls.filter(effectiveness_rating=effectiveness_rating)
    
    # Calculate statistics
    total_controls = controls.count()
    implemented_controls = controls.filter(implementation_status='fully_implemented').count()
    effective_controls = controls.filter(effectiveness_rating='effective').count()
    highly_effective_controls = controls.filter(effectiveness_rating='highly_effective').count()
    
    filters = {
        'process': process,
        'control_type': control_type,
        'implementation_status': implementation_status,
        'effectiveness_rating': effectiveness_rating
    }
    
    html_string = render_to_string('reports/cobit_control_effectiveness.html', {
        'organization': org,
        'controls': controls,
        'total_controls': total_controls,
        'implemented_controls': implemented_controls,
        'effective_controls': effective_controls,
        'highly_effective_controls': highly_effective_controls,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_cobit_control_effectiveness.pdf"'
    return response

# NIST Reports
def nist_function_summary_pdf(request):
    """Generate NIST function summary report"""
    from risk.models import NISTFunction, NISTCategory, NISTSubcategory, NISTImplementation
    
    org = request.tenant
    functions = NISTFunction.objects.filter(organization=org)
    
    # Apply filters
    function_code = request.GET.get('function_code')
    if function_code:
        functions = functions.filter(function_code=function_code)
    
    # Calculate statistics
    total_categories = NISTCategory.objects.filter(function__organization=org).count()
    total_subcategories = NISTSubcategory.objects.filter(category__function__organization=org).count()
    total_implementations = NISTImplementation.objects.filter(subcategory__category__function__organization=org).count()
    
    filters = {'function_code': function_code}
    
    html_string = render_to_string('reports/nist_function_summary.html', {
        'organization': org,
        'functions': functions,
        'total_categories': total_categories,
        'total_subcategories': total_subcategories,
        'total_implementations': total_implementations,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_nist_function_summary.pdf"'
    return response

def nist_implementation_status_pdf(request):
    """Generate NIST implementation status report"""
    from risk.models import NISTImplementation
    
    org = request.tenant
    implementations = NISTImplementation.objects.filter(subcategory__category__function__organization=org)
    
    # Apply filters
    subcategory = request.GET.get('subcategory')
    current_maturity = request.GET.get('current_maturity')
    target_maturity = request.GET.get('target_maturity')
    implementation_status = request.GET.get('implementation_status')
    
    if subcategory:
        implementations = implementations.filter(subcategory__subcategory_code__icontains=subcategory)
    if current_maturity:
        implementations = implementations.filter(current_maturity=current_maturity)
    if target_maturity:
        implementations = implementations.filter(target_maturity=target_maturity)
    if implementation_status:
        implementations = implementations.filter(implementation_status=implementation_status)
    
    # Calculate statistics
    avg_current_maturity = implementations.aggregate(avg=models.Avg('current_maturity'))['avg'] or 0
    avg_target_maturity = implementations.aggregate(avg=models.Avg('target_maturity'))['avg'] or 0
    maturity_gap = avg_target_maturity - avg_current_maturity
    fully_implemented = implementations.filter(implementation_status='fully_implemented').count()
    
    filters = {
        'subcategory': subcategory,
        'current_maturity': current_maturity,
        'target_maturity': target_maturity,
        'implementation_status': implementation_status
    }
    
    html_string = render_to_string('reports/nist_implementation_status.html', {
        'organization': org,
        'implementations': implementations,
        'avg_current_maturity': avg_current_maturity,
        'avg_target_maturity': avg_target_maturity,
        'maturity_gap': maturity_gap,
        'fully_implemented': fully_implemented,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_nist_implementation_status.pdf"'
    return response

def nist_incident_register_pdf(request):
    """Generate NIST incident register report"""
    from risk.models import NISTIncident
    
    org = request.tenant
    incidents = NISTIncident.objects.filter(organization=org)
    
    # Apply filters
    incident_type = request.GET.get('incident_type')
    severity = request.GET.get('severity')
    status = request.GET.get('status')
    
    if incident_type:
        incidents = incidents.filter(incident_type=incident_type)
    if severity:
        incidents = incidents.filter(severity=severity)
    if status:
        incidents = incidents.filter(status=status)
    
    # Calculate statistics
    total_incidents = incidents.count()
    open_incidents = incidents.exclude(status__in=['closed', 'recovered']).count()
    high_critical_incidents = incidents.filter(severity__in=['high', 'critical']).count()
    
    # Calculate average resolution time
    resolved_incidents = incidents.filter(resolved_date__isnull=False)
    avg_resolution_time = 0
    if resolved_incidents.exists():
        total_days = sum((incident.resolved_date - incident.detected_date).days for incident in resolved_incidents)
        avg_resolution_time = total_days / resolved_incidents.count()
    
    filters = {
        'incident_type': incident_type,
        'severity': severity,
        'status': status
    }
    
    html_string = render_to_string('reports/nist_incident_register.html', {
        'organization': org,
        'incidents': incidents,
        'total_incidents': total_incidents,
        'open_incidents': open_incidents,
        'high_critical_incidents': high_critical_incidents,
        'avg_resolution_time': avg_resolution_time,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_nist_incident_register.pdf"'
    return response

def nist_threat_analysis_pdf(request):
    """Generate NIST threat analysis report"""
    from risk.models import NISTThreat
    
    org = request.tenant
    threats = NISTThreat.objects.filter(organization=org)
    
    # Apply filters
    threat_type = request.GET.get('threat_type')
    severity = request.GET.get('severity')
    likelihood = request.GET.get('likelihood')
    
    if threat_type:
        threats = threats.filter(threat_type=threat_type)
    if severity:
        threats = threats.filter(severity=severity)
    if likelihood:
        threats = threats.filter(likelihood=likelihood)
    
    # Calculate statistics
    total_threats = threats.count()
    high_critical_threats = threats.filter(severity__in=['high', 'critical']).count()
    likely_certain_threats = threats.filter(likelihood__in=['likely', 'certain']).count()
    adversarial_threats = threats.filter(threat_type='adversarial').count()
    
    filters = {
        'threat_type': threat_type,
        'severity': severity,
        'likelihood': likelihood
    }
    
    html_string = render_to_string('reports/nist_threat_analysis.html', {
        'organization': org,
        'threats': threats,
        'total_threats': total_threats,
        'high_critical_threats': high_critical_threats,
        'likely_certain_threats': likely_certain_threats,
        'adversarial_threats': adversarial_threats,
        'filters': filters,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_nist_threat_analysis.pdf"'
    return response

# ─── OBJECTIVE REPORTS (RISK APP) ─────────────────────────────────────────────
def objective_list_pdf(request):
    org = request.organization
    q = request.GET.get('q')
    status = request.GET.get('status')
    objectives = Objective.objects.filter(organization=org)
    if q:
        objectives = objectives.filter(
            Q(name__icontains=q) | Q(code__icontains=q) | Q(origin_source__icontains=q)
        )
    if status:
        objectives = objectives.filter(status=status)

    # Aggregate simple insights per objective
    rows = []
    for obj in objectives:
        risks_qs = obj.risks.all()
        rows.append({
            'objective': obj,
            'risk_count': risks_qs.count(),
            'open_risks': risks_qs.filter(status='open').count(),
            'in_progress_risks': risks_qs.filter(status='in-progress').count(),
            'closed_risks': risks_qs.filter(status='closed').count(),
        })

    html_string = render_to_string('reports/objective_list.html', {
        'organization': org,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'rows': rows,
        'filters': {'q': q, 'status': status},
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_objective_list.pdf"'
    return response

def objective_detailed_pdf(request):
    org = request.organization
    # Filters: objective status and search
    q = request.GET.get('q')
    status = request.GET.get('status')
    objectives = Objective.objects.filter(organization=org)
    if q:
        objectives = objectives.filter(
            Q(name__icontains=q) | Q(code__icontains=q) | Q(origin_source__icontains=q)
        )
    if status:
        objectives = objectives.filter(status=status)

    # Build context with risks per objective and what is being done (action plan fields)
    data = []
    for obj in objectives:
        risks = obj.risks.all().order_by('-updated_at')
        risk_rows = []
        for r in risks:
            risk_rows.append({
                'risk': r,
                'action_plan': r.action_plan,
                'action_plan_status': r.action_plan_status,
                'action_owner': r.action_owner,
                'action_due_date': r.action_due_date,
                'control_status': r.control_status,
                'control_rating': r.control_rating,
            })
        data.append({'objective': obj, 'risk_rows': risk_rows})

    html_string = render_to_string('reports/objective_detailed.html', {
        'organization': org,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'data': data,
        'filters': {'q': q, 'status': status},
    })
    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(
        stylesheets=[CSS(string='@page { size: A4; margin: 1.5cm }')]
    )
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{org.code}_objective_detailed.pdf"'
    return response


# AI Governance Reports
def ai_governance_dashboard_pdf(request):
    """Generate AI Governance Dashboard PDF report."""
    from ai_governance.reports import ai_governance_dashboard_pdf as ai_dashboard_pdf
    return ai_dashboard_pdf(request)


def ai_governance_test_run_details_pdf(request):
    """Generate detailed test run report."""
    from ai_governance.reports import test_run_details_pdf as ai_test_run_pdf
    return ai_test_run_pdf(request)


def ai_governance_compliance_matrix_pdf(request):
    """Generate compliance matrix report showing test mappings to frameworks."""
    from ai_governance.reports import compliance_matrix_pdf as ai_compliance_pdf
    return ai_compliance_pdf(request)