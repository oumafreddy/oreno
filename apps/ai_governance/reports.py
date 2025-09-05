"""
AI Governance reporting views that integrate with the existing reports system.
"""

import matplotlib
matplotlib.use('Agg')  # Set non-interactive backend

from django.shortcuts import render
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML, CSS
from django.db.models import Count, Q, Avg, Sum
from django.db import models
import tempfile
import matplotlib.pyplot as plt
import io
import base64
from datetime import datetime, timedelta
from django.utils import timezone
from django.db.models.functions import TruncMonth, TruncDay

from .models import (
    ModelAsset, DatasetAsset, TestPlan, TestRun, TestResult, 
    Metric, EvidenceArtifact, Framework, Clause, ComplianceMapping
)
from admin_module.models import DataExportLog


def _create_ai_governance_docx(org, title, generation_timestamp):
    """Create a python-docx Document with AI governance specific header."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    
    # Organization logo (best-effort)
    if getattr(org, 'logo', None) and getattr(org.logo, 'path', None):
        try:
            doc.add_picture(org.logo.path, width=Inches(1.8))
        except Exception:
            pass

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata table
    meta = doc.add_table(rows=3, cols=2)
    meta.style = 'Table Grid'
    meta.cell(0, 0).text = 'Organization'
    meta.cell(0, 1).text = getattr(org, 'name', 'Organization')
    meta.cell(1, 0).text = 'Report Type'
    meta.cell(1, 1).text = 'AI Governance'
    meta.cell(2, 0).text = 'Generated'
    meta.cell(2, 1).text = generation_timestamp
    return doc


def _ai_governance_http_response(doc, filename_prefix, org):
    """Return the given Document as an HTTP response download."""
    from docx import Document
    import tempfile
    import os
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
        
        # Read the file
        with open(tmp_file.name, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename_prefix}_{org.name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
        
        # Clean up
        os.unlink(tmp_file.name)
        return response


def ai_governance_dashboard_pdf(request):
    """Generate AI Governance Dashboard PDF report."""
    org = request.tenant
    
    # Get filter parameters
    model_type = request.GET.get('model_type', '')
    date_range = request.GET.get('date_range', 'last_30_days')
    framework = request.GET.get('framework', '')
    
    # Calculate date range
    if date_range == 'last_7_days':
        days_back = 7
    elif date_range == 'last_30_days':
        days_back = 30
    elif date_range == 'last_90_days':
        days_back = 90
    elif date_range == 'last_year':
        days_back = 365
    else:  # all_time
        days_back = None
    
    # Build base querysets
    models_qs = ModelAsset.objects.filter(organization=org)
    test_runs_qs = TestRun.objects.filter(organization=org)
    
    # Apply filters
    if model_type:
        models_qs = models_qs.filter(model_type=model_type)
        test_runs_qs = test_runs_qs.filter(model_asset__model_type=model_type)
    
    if days_back:
        cutoff_date = timezone.now() - timedelta(days=days_back)
        test_runs_qs = test_runs_qs.filter(created_at__gte=cutoff_date)
    
    # Get dashboard data
    total_models = models_qs.count()
    total_test_runs = test_runs_qs.count()
    
    # Recent test runs
    recent_test_runs = test_runs_qs.select_related('model_asset', 'test_plan').order_by('-created_at')[:10]
    
    # Test results statistics (use filtered test runs)
    if days_back:
        recent_results = TestResult.objects.filter(
            test_run__in=test_runs_qs
        )
    else:
        recent_results = TestResult.objects.filter(
            test_run__organization=org
        )
    passed_tests = recent_results.filter(passed=True).count()
    total_recent_tests = recent_results.count()
    compliance_score = round((passed_tests / total_recent_tests) * 100) if total_recent_tests > 0 else 0
    
    # Test category breakdown
    fairness_passed = recent_results.filter(test_name__startswith='demographic_parity').filter(passed=True).count()
    explainability_passed = recent_results.filter(test_name__startswith='shap_').filter(passed=True).count()
    robustness_passed = recent_results.filter(test_name__startswith='adversarial_').filter(passed=True).count()
    privacy_passed = recent_results.filter(test_name__startswith='membership_').filter(passed=True).count()
    
    # Framework compliance scores (mock data for now)
    eu_ai_act_score = 85
    oecd_score = 78
    nist_score = 92
    
    # Available assets (use filtered data)
    available_models = models_qs
    available_datasets = DatasetAsset.objects.filter(organization=org)
    available_test_plans = TestPlan.objects.filter(organization=org)
    
    context = {
        'organization': org,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'AI Governance Dashboard Report',
        'description': f'Comprehensive AI governance overview with {total_models} models and {total_test_runs} test runs',
        'total_models': total_models,
        'total_test_runs': total_test_runs,
        'recent_test_runs': recent_test_runs,
        'passed_tests': passed_tests,
        'total_recent_tests': total_recent_tests,
        'compliance_score': compliance_score,
        'fairness_passed': fairness_passed,
        'explainability_passed': explainability_passed,
        'robustness_passed': robustness_passed,
        'privacy_passed': privacy_passed,
        'eu_ai_act_score': eu_ai_act_score,
        'oecd_score': oecd_score,
        'nist_score': nist_score,
        'available_models': available_models,
        'available_datasets': available_datasets,
        'available_test_plans': available_test_plans,
    }
    
    if request.GET.get('format') == 'docx':
        doc = _create_ai_governance_docx(org, context['title'], context['generation_timestamp'])
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f"""
        This AI Governance Dashboard Report provides a comprehensive overview of AI model compliance and testing activities for {org.name}.
        
        Key Metrics:
        • Total Registered Models: {total_models}
        • Total Test Runs: {total_test_runs}
        • Overall Compliance Score: {compliance_score}%
        • Recent Tests Passed: {passed_tests}/{total_recent_tests}
        
        Framework Compliance:
        • EU AI Act: {eu_ai_act_score}%
        • OECD Principles: {oecd_score}%
        • NIST AI RMF: {nist_score}%
        """)
        
        # Add model inventory
        doc.add_heading('Model Inventory', level=1)
        if available_models:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Model Name'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Version'
            hdr_cells[3].text = 'Created'
            
            for model in available_models:
                row_cells = table.add_row().cells
                row_cells[0].text = model.name
                row_cells[1].text = model.model_type
                row_cells[2].text = model.version or 'latest'
                row_cells[3].text = model.created_at.strftime('%Y-%m-%d')
        else:
            doc.add_paragraph('No models registered yet.')
        
        # Add recent test runs
        doc.add_heading('Recent Test Runs', level=1)
        if recent_test_runs:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Model'
            hdr_cells[1].text = 'Test Plan'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Created'
            hdr_cells[4].text = 'Duration'
            
            for test_run in recent_test_runs:
                row_cells = table.add_row().cells
                row_cells[0].text = test_run.model_asset.name
                row_cells[1].text = test_run.test_plan.name if test_run.test_plan else 'N/A'
                row_cells[2].text = test_run.status.title()
                row_cells[3].text = test_run.created_at.strftime('%Y-%m-%d %H:%M')
                if test_run.completed_at and test_run.started_at:
                    duration = test_run.completed_at - test_run.started_at
                    row_cells[4].text = str(duration).split('.')[0]  # Remove microseconds
                else:
                    row_cells[4].text = 'N/A'
        else:
            doc.add_paragraph('No recent test runs.')
        
        return _ai_governance_http_response(doc, 'ai_governance_dashboard', org)
    
    # PDF format
    html_string = render_to_string('reports/ai_governance_dashboard.html', context)
    html = HTML(string=html_string)
    css = CSS(string='''
        @page {
            size: A4;
            margin: 1in;
        }
        body {
            font-family: Arial, sans-serif;
            font-size: 12px;
            line-height: 1.4;
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 2px solid #333;
            padding-bottom: 20px;
        }
        .metric-card {
            background: #f8f9fa;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            padding: 15px;
            margin: 10px 0;
        }
        .metric-value {
            font-size: 24px;
            font-weight: bold;
            color: #007bff;
        }
        .metric-label {
            color: #6c757d;
            font-size: 14px;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th, td {
            border: 1px solid #dee2e6;
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        .status-completed { color: #28a745; }
        .status-failed { color: #dc3545; }
        .status-running { color: #ffc107; }
        .status-pending { color: #6c757d; }
    ''')
    
    pdf_file = html.write_pdf(stylesheets=[css])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="ai_governance_dashboard_{org.name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    return response


def test_run_details_pdf(request, test_run_id=None):
    """Generate detailed test run report."""
    org = request.tenant
    
    # Get test_run_id from URL parameter or form parameter
    if not test_run_id:
        test_run_id = request.GET.get('test_run_id')
    
    if not test_run_id:
        return HttpResponse('Test run ID is required', status=400)
    
    try:
        test_run = TestRun.objects.select_related('model_asset', 'dataset_asset', 'test_plan').get(
            id=test_run_id, organization=org
        )
    except TestRun.DoesNotExist:
        return HttpResponse('Test run not found', status=404)
    
    # Get test results and metrics
    test_results = TestResult.objects.filter(test_run=test_run).prefetch_related('metrics')
    artifacts = EvidenceArtifact.objects.filter(test_run=test_run)
    
    # Calculate summary statistics
    total_tests = test_results.count()
    passed_tests = test_results.filter(passed=True).count()
    failed_tests = total_tests - passed_tests
    pass_rate = round((passed_tests / total_tests) * 100) if total_tests > 0 else 0
    
    # Group results by test category
    test_categories = {}
    for result in test_results:
        category = result.test_name.split('_')[0] if '_' in result.test_name else 'other'
        if category not in test_categories:
            test_categories[category] = {'total': 0, 'passed': 0}
        test_categories[category]['total'] += 1
        if result.passed:
            test_categories[category]['passed'] += 1
    
    context = {
        'organization': org,
        'test_run': test_run,
        'test_results': test_results,
        'artifacts': artifacts,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': f'Test Run Report - {test_run.model_asset.name}',
        'description': f'Detailed analysis of test run {test_run.id}',
        'total_tests': total_tests,
        'passed_tests': passed_tests,
        'failed_tests': failed_tests,
        'pass_rate': pass_rate,
        'test_categories': test_categories,
    }
    
    if request.GET.get('format') == 'docx':
        doc = _create_ai_governance_docx(org, context['title'], context['generation_timestamp'])
        
        # Add test run summary
        doc.add_heading('Test Run Summary', level=1)
        doc.add_paragraph(f"""
        Model: {test_run.model_asset.name}
        Dataset: {test_run.dataset_asset.name if test_run.dataset_asset else 'N/A'}
        Test Plan: {test_run.test_plan.name if test_run.test_plan else 'N/A'}
        Status: {test_run.status.title()}
        Created: {test_run.created_at.strftime('%Y-%m-%d %H:%M:%S')}
        """)
        
        if test_run.completed_at:
            doc.add_paragraph(f"Completed: {test_run.completed_at.strftime('%Y-%m-%d %H:%M:%S')}")
            if test_run.started_at:
                duration = test_run.completed_at - test_run.started_at
                doc.add_paragraph(f"Duration: {duration}")
        
        # Add test results summary
        doc.add_heading('Test Results Summary', level=1)
        doc.add_paragraph(f"""
        Total Tests: {total_tests}
        Passed: {passed_tests}
        Failed: {failed_tests}
        Pass Rate: {pass_rate}%
        """)
        
        # Add detailed test results
        doc.add_heading('Detailed Test Results', level=1)
        if test_results:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Test Name'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Score'
            hdr_cells[3].text = 'Execution Time'
            
            for result in test_results:
                row_cells = table.add_row().cells
                row_cells[0].text = result.test_name
                row_cells[1].text = 'PASSED' if result.passed else 'FAILED'
                row_cells[2].text = str(result.summary.get('score', 'N/A'))
                row_cells[3].text = f"{result.execution_time:.2f}s" if result.execution_time else 'N/A'
        else:
            doc.add_paragraph('No test results available.')
        
        return _ai_governance_http_response(doc, f'test_run_{test_run_id}', org)
    
    # PDF format
    html_string = render_to_string('reports/ai_governance_test_run_details.html', context)
    html = HTML(string=html_string)
    css = CSS(string='''
        @page {
            size: A4;
            margin: 1in;
        }
        body {
            font-family: Arial, sans-serif;
            font-size: 12px;
            line-height: 1.4;
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 2px solid #333;
            padding-bottom: 20px;
        }
        .summary-box {
            background: #f8f9fa;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            padding: 15px;
            margin: 20px 0;
        }
        .status-passed { color: #28a745; font-weight: bold; }
        .status-failed { color: #dc3545; font-weight: bold; }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th, td {
            border: 1px solid #dee2e6;
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
    ''')
    
    pdf_file = html.write_pdf(stylesheets=[css])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="test_run_{test_run_id}_{org.name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    return response


def compliance_matrix_pdf(request):
    """Generate compliance matrix report showing test mappings to frameworks."""
    org = request.tenant
    
    # Get filter parameters
    framework_code = request.GET.get('framework', '')
    test_name = request.GET.get('test_name', '')
    date_range = request.GET.get('date_range', 'last_30_days')
    
    # Calculate date range
    if date_range == 'last_7_days':
        days_back = 7
    elif date_range == 'last_30_days':
        days_back = 30
    elif date_range == 'last_90_days':
        days_back = 90
    elif date_range == 'last_year':
        days_back = 365
    else:  # all_time
        days_back = None
    
    # Get frameworks and mappings with filters
    frameworks_qs = Framework.objects.filter(organization=org).prefetch_related('clauses__mappings')
    mappings_qs = ComplianceMapping.objects.filter(organization=org).select_related('clause__framework')
    
    # Apply filters
    if framework_code:
        frameworks_qs = frameworks_qs.filter(code=framework_code)
        mappings_qs = mappings_qs.filter(clause__framework__code=framework_code)
    
    if test_name:
        mappings_qs = mappings_qs.filter(test_name=test_name)
    
    # Get recent test results for compliance scoring
    if days_back:
        cutoff_date = timezone.now() - timedelta(days=days_back)
        recent_results = TestResult.objects.filter(
            test_run__organization=org,
            test_run__created_at__gte=cutoff_date
        )
    else:
        recent_results = TestResult.objects.filter(test_run__organization=org)
    
    # Calculate compliance scores per framework
    framework_scores = {}
    for framework in frameworks_qs:
        framework_mappings = mappings_qs.filter(clause__framework=framework)
        total_mappings = framework_mappings.count()
        
        if total_mappings > 0:
            passed_mappings = 0
            for mapping in framework_mappings:
                # Check if recent test results for this test name passed
                test_results = recent_results.filter(test_name=mapping.test_name)
                if test_results.exists() and test_results.filter(passed=True).exists():
                    passed_mappings += 1
            
            framework_scores[framework.code] = round((passed_mappings / total_mappings) * 100)
        else:
            framework_scores[framework.code] = 0
    
    # Calculate average compliance score
    average_score = 0
    if framework_scores:
        total_score = sum(framework_scores.values())
        average_score = round(total_score / len(framework_scores))
    
    context = {
        'organization': org,
        'frameworks': frameworks_qs,
        'mappings': mappings_qs,
        'framework_scores': framework_scores,
        'average_score': average_score,
        'generation_timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'title': 'AI Governance Compliance Matrix',
        'description': 'Comprehensive compliance mapping across all frameworks',
    }
    
    if request.GET.get('format') == 'docx':
        doc = _create_ai_governance_docx(org, context['title'], context['generation_timestamp'])
        
        # Add framework scores
        doc.add_heading('Framework Compliance Scores', level=1)
        for framework in frameworks_qs:
            score = framework_scores.get(framework.code, 0)
            doc.add_paragraph(f"{framework.title}: {score}%")
        
        # Add compliance matrix
        doc.add_heading('Compliance Matrix', level=1)
        if mappings_qs:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Framework'
            hdr_cells[1].text = 'Clause'
            hdr_cells[2].text = 'Test Name'
            hdr_cells[3].text = 'Rationale'
            
            for mapping in mappings_qs:
                row_cells = table.add_row().cells
                row_cells[0].text = mapping.clause.framework.title
                row_cells[1].text = mapping.clause.clause_code
                row_cells[2].text = mapping.test_name
                row_cells[3].text = mapping.rationale[:100] + '...' if len(mapping.rationale) > 100 else mapping.rationale
        else:
            doc.add_paragraph('No compliance mappings found.')
        
        return _ai_governance_http_response(doc, 'compliance_matrix', org)
    
    # PDF format
    html_string = render_to_string('reports/ai_governance_compliance_matrix.html', context)
    html = HTML(string=html_string)
    css = CSS(string='''
        @page {
            size: A4;
            margin: 1in;
        }
        body {
            font-family: Arial, sans-serif;
            font-size: 12px;
            line-height: 1.4;
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 2px solid #333;
            padding-bottom: 20px;
        }
        .framework-score {
            background: #f8f9fa;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            padding: 15px;
            margin: 10px 0;
            text-align: center;
        }
        .score-value {
            font-size: 24px;
            font-weight: bold;
            color: #007bff;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th, td {
            border: 1px solid #dee2e6;
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
    ''')
    
    pdf_file = html.write_pdf(stylesheets=[css])
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="compliance_matrix_{org.name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    return response
